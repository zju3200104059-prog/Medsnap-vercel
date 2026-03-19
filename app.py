# -*- coding: utf-8 -*-
"""
临床科研病历AI识别与结构化提取工具 - 主程序（多角色版）
Web服务监听: 0.0.0.0:7860
支持角色: 临床医生 / 护士 / 临床科研人员
"""

import os
import json
import uuid
import base64
import tempfile
import sqlite3
import shutil
import re
from datetime import datetime
from flask import Flask, render_template, request, send_file, jsonify
from openai import OpenAI
from PIL import Image, ImageFilter, ImageEnhance, ImageOps
import numpy as np
import pandas as pd
from desensitizer import desensitize_text, desensitize_structured_data

HAS_PYMUPDF = False
try:
    import fitz
    HAS_PYMUPDF = True
except ImportError:
    pass

HAS_DASHSCOPE = False
try:
    import dashscope
    from dashscope.audio.asr import Recognition
    HAS_DASHSCOPE = True
except ImportError:
    pass

HAS_PYDUB = False
try:
    from pydub import AudioSegment
    HAS_PYDUB = True
except ImportError:
    pass

HAS_TESSERACT = False
try:
    import pytesseract
    HAS_TESSERACT = True
    # Windows下Tesseract默认安装路径
    import platform
    if platform.system() == 'Windows':
        _tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        for _tp in _tesseract_paths:
            if os.path.exists(_tp):
                pytesseract.pytesseract.tesseract_cmd = _tp
                break
except ImportError:
    pass

# ========== Flask 应用初始化 ==========
app = Flask(__name__)
app.secret_key = uuid.uuid4().hex
app.config['JSON_AS_ASCII'] = False
try:
    app.json.ensure_ascii = False
except (AttributeError, TypeError):
    pass

UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "medical_ocr_uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_DIR
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

from statistics_routes import stats_bp
app.register_blueprint(stats_bp)

from research_routes import research_bp
app.register_blueprint(research_bp)

DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'local_medical_data.db')

# ========== 多模态大模型配置 ==========
API_KEY = os.environ.get("MODELSCOPE_API_KEY", "")
if not API_KEY:
    print("[WARN] 环境变量 MODELSCOPE_API_KEY 未设置，AI识别功能将不可用")
client = OpenAI(
    base_url='https://api-inference.modelscope.cn/v1',
    api_key=API_KEY
)
MODEL_NAME = "Qwen/Qwen3-VL-235B-A22B-Instruct"

# ========== 语音识别配置 ==========
DASHSCOPE_API_KEY = os.environ.get("DASHSCOPE_API_KEY", "")
if not DASHSCOPE_API_KEY:
    print("[WARN] 环境变量 DASHSCOPE_API_KEY 未设置，语音识别功能将不可用")
if HAS_DASHSCOPE and DASHSCOPE_API_KEY:
    dashscope.api_key = DASHSCOPE_API_KEY

ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'bmp', 'tiff', 'pdf'}
ALLOWED_AUDIO_EXTENSIONS = {'wav', 'mp3', 'aac', 'amr', 'opus', 'm4a', 'flac'}
ALLOWED_TEXT_EXTENSIONS = {'txt', 'doc', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def is_audio_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_AUDIO_EXTENSIONS

def is_text_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_TEXT_EXTENSIONS


# ========== 科室与模板配置 ==========
DEPARTMENT_CONFIGS = {
    'cardiology':  {'name': '心内科',   'color': '#dc2626'},
    'neurology':   {'name': '神经内科', 'color': '#7c3aed'},
    'surgery':     {'name': '外科',     'color': '#0891b2'},
    'pediatrics':  {'name': '儿科',     'color': '#f59e0b'},
    'obstetrics':  {'name': '妇产科',   'color': '#ec4899'},
    'emergency':   {'name': '急诊科',   'color': '#ef4444'},
    'general':     {'name': '通用',     'color': '#64748b'},
}

# 向后兼容：旧角色ID映射到通用科室
LEGACY_ROLE_MAP = {
    'diagnosis': 'general', 'nursing': 'general', 'other': 'general',
    'doctor': 'general', 'nurse': 'general', 'researcher': 'general',
}

# 保留旧名供内部兼容
CATEGORY_CONFIGS = DEPARTMENT_CONFIGS

# ========== 内置模板Prompt定义 ==========

PROMPT_DOCTOR_MEDICAL_RECORD = """你是临床医生数据提取专家。请仔细识别该病历图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记、代码块标记或多余文字。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "", "床号": "", "住院号": ""
  },
  "chief_complaint": "",
  "present_illness": "",
  "past_history": "",
  "personal_history": "",
  "family_history": "",
  "physical_exam": {
    "体温": null, "脉搏": null, "呼吸": null, "血压": "",
    "一般情况": "", "专科检查": ""
  },
  "diagnosis": [
    {"诊断名称": "", "ICD10编码": ""}
  ],
  "treatment_plan": {
    "药物治疗": "", "手术治疗": "", "其他治疗": ""
  },
  "surgery_record": "",
  "discharge_summary": "",
  "confidence": {}
}

## 提取规则
1. 年龄提取纯数字（如"56岁"→56）。
2. 诊断需尽量识别ICD-10编码（如 E11.9 2型糖尿病、I10 高血压病）。
3. 治疗方案区分药物治疗、手术治疗、其他治疗。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。未识别到的字段置信度为0。
5. 手写体请尽力识别，无法确认的字用?标记。

只输出JSON，不要输出任何其他内容。"""

PROMPT_DOCTOR_LAB_RESULTS = """你是临床检验数据提取专家。请仔细识别该检查检验结果图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记、代码块标记或多余文字。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "", "住院号": ""
  },
  "report_info": {
    "报告类型": "", "检查日期": "", "报告日期": ""
  },
  "lab_tests": [
    {
      "项目名称": "", "英文缩写": "", "数值": null,
      "单位": "", "参考范围": "", "异常标注": ""
    }
  ],
  "confidence": {}
}

## 提取规则
1. 识别常见医学缩写：WBC、RBC、PLT、Hb、ALT、AST、Cr、BUN、GLU、TC、TG、HDL-C、LDL-C、UA、CRP、ESR、HbA1c、TSH、FT3、FT4、Na、K、Ca、Cl、PT、APTT、INR、D-Dimer、AFP、CEA等。
2. 异常标注用"↑"(偏高)/"↓"(偏低)/"正常"。
3. 数值标准化为数字格式。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。

## 医疗专业词汇参考
- 血常规：WBC、RBC、PLT、Hb、HCT、MCV、MCH、MCHC、RDW
- 肝功能：ALT、AST、GGT、ALP、TBIL、DBIL、TP、ALB
- 肾功能：Cr、BUN、UA、Cys-C、eGFR
- 血脂：TC、TG、HDL-C、LDL-C、ApoA1、ApoB
- 血糖：GLU、FPG、2hPG、HbA1c、OGTT
- 凝血：PT、APTT、TT、FIB、INR、D-Dimer
- 电解质：Na、K、Ca、Cl、Mg、P
- 炎症指标：CRP、PCT、ESR、IL-6
- 肿瘤标志物：AFP、CEA、CA125、CA199、CA153、PSA

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_ADMISSION = """你是护理评估数据提取专家。请仔细识别该入院护理评估表图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "", "床号": "", "住院号": "", "入院日期": ""
  },
  "vital_signs": {
    "体温": null, "脉搏": null, "呼吸": null, "血压_收缩压": null, "血压_舒张压": null,
    "血氧饱和度": null, "身高_cm": null, "体重_kg": null
  },
  "assessment": {
    "意识状态": "", "精神状态": "", "皮肤完整性": "", "皮肤异常描述": "",
    "营养状况": "", "饮食类型": "", "排便情况": "", "排尿情况": "",
    "睡眠情况": "", "活动能力": "", "自理能力初筛": "",
    "跌倒风险初筛": "", "压疮风险初筛": "", "疼痛评分": null,
    "过敏史": "", "特殊用药": ""
  },
  "nursing_diagnosis": [],
  "nursing_plan": "",
  "confidence": {}
}

## 提取规则
1. 生命体征数值标准化为纯数字。血压格式拆分为收缩压和舒张压。
2. 意识状态：清醒/嗜睡/昏睡/浅昏迷/深昏迷。
3. 自理能力/跌倒/压疮风险初筛：识别勾选框或评分。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_BARTHEL = """你是护理评估数据提取专家。请仔细识别该Barthel自理能力指数量表图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "barthel_items": [
    {"项目": "进食", "评分": null, "评分标准": "10=自理 5=需部分帮助 0=完全依赖", "备注": ""},
    {"项目": "洗澡", "评分": null, "评分标准": "5=自理 0=需帮助", "备注": ""},
    {"项目": "修饰", "评分": null, "评分标准": "5=自理 0=需帮助", "备注": ""},
    {"项目": "穿衣", "评分": null, "评分标准": "10=自理 5=需部分帮助 0=完全依赖", "备注": ""},
    {"项目": "控制大便", "评分": null, "评分标准": "10=可控 5=偶有失禁 0=失禁", "备注": ""},
    {"项目": "控制小便", "评分": null, "评分标准": "10=可控 5=偶有失禁 0=失禁", "备注": ""},
    {"项目": "如厕", "评分": null, "评分标准": "10=自理 5=需部分帮助 0=完全依赖", "备注": ""},
    {"项目": "床椅转移", "评分": null, "评分标准": "15=自理 10=少量帮助 5=较大帮助 0=完全依赖", "备注": ""},
    {"项目": "平地行走", "评分": null, "评分标准": "15=自行45m 10=在帮助下45m 5=轮椅45m 0=不能", "备注": ""},
    {"项目": "上下楼梯", "评分": null, "评分标准": "10=自理 5=需帮助 0=不能", "备注": ""}
  ],
  "total_score": null,
  "dependency_level": "",
  "confidence": {}
}

## 提取规则
1. 评分只提取数字。
2. 总分范围0-100。依赖等级判定：100=自理、61-99=轻度依赖、41-60=中度依赖、≤40=重度依赖。
3. 如图中有勾选标记，按勾选对应的分值提取。
4. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_MORSE = """你是护理评估数据提取专家。请仔细识别该Morse跌倒风险评估量表图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "morse_items": [
    {"项目": "跌倒史", "评分": null, "评分标准": "是=25 否=0", "选项": ""},
    {"项目": "继发诊断", "评分": null, "评分标准": "是=15 否=0", "选项": ""},
    {"项目": "步行辅助", "评分": null, "评分标准": "卧床/护士协助=0 拐杖/助行器/轮椅=15 扶家具行走=30", "选项": ""},
    {"项目": "静脉输液/肝素锁", "评分": null, "评分标准": "是=20 否=0", "选项": ""},
    {"项目": "步态", "评分": null, "评分标准": "正常/卧床/不能活动=0 虚弱=10 损伤=20", "选项": ""},
    {"项目": "认知状态", "评分": null, "评分标准": "能正确认识自身活动能力=0 高估/忘记限制=15", "选项": ""}
  ],
  "total_score": null,
  "risk_level": "",
  "confidence": {}
}

## 提取规则
1. 评分只提取数字。
2. 风险判定：0-24=低风险、25-44=中风险、≥45=高风险。
3. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_BRADEN = """你是护理评估数据提取专家。请仔细识别该Braden压疮风险评估量表图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "braden_items": [
    {"项目": "感知能力", "评分": null, "评分标准": "1=完全受限 2=非常受限 3=轻度受限 4=未受损", "备注": ""},
    {"项目": "潮湿程度", "评分": null, "评分标准": "1=持续潮湿 2=非常潮湿 3=偶尔潮湿 4=很少潮湿", "备注": ""},
    {"项目": "活动能力", "评分": null, "评分标准": "1=卧床 2=坐椅 3=偶尔步行 4=经常步行", "备注": ""},
    {"项目": "移动能力", "评分": null, "评分标准": "1=完全不能 2=严重受限 3=轻度受限 4=不受限", "备注": ""},
    {"项目": "营养摄取", "评分": null, "评分标准": "1=非常差 2=可能不足 3=足够 4=良好", "备注": ""},
    {"项目": "摩擦力和剪切力", "评分": null, "评分标准": "1=存在问题 2=潜在问题 3=不存在问题", "备注": ""}
  ],
  "total_score": null,
  "risk_level": "",
  "confidence": {}
}

## 提取规则
1. 评分只提取数字（1-4分，摩擦力1-3分）。
2. 总分6-23分。风险判定：≤9=极高风险、10-12=高风险、13-14=中度风险、15-18=低风险、≥19=无风险。
3. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_PAIN = """你是护理评估数据提取专家。请仔细识别该疼痛评估记录图片中的所有内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "评估日期": "", "评估人": ""
  },
  "pain_assessment": {
    "疼痛部位": "",
    "疼痛性质": "",
    "NRS评分": null,
    "VAS评分": null,
    "疼痛频率": "",
    "持续时间": "",
    "加重因素": "",
    "缓解因素": "",
    "对睡眠影响": "",
    "对日常活动影响": "",
    "当前镇痛措施": "",
    "镇痛效果": ""
  },
  "confidence": {}
}

## 提取规则
1. NRS评分0-10（0=无痛，10=最剧烈疼痛）。VAS评分0-10。
2. 疼痛性质：锐痛/钝痛/刺痛/胀痛/灼痛/绞痛等。
3. confidence字段：对每个项目给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_NURSE_RECORD = """你是护理记录数据提取专家。请仔细识别该护理记录单图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "科室": "", "床号": "", "记录日期": ""
  },
  "vital_signs_records": [
    {"时间": "", "体温": null, "脉搏": null, "呼吸": null, "血压": "", "血氧": null, "备注": ""}
  ],
  "medication_execution": [
    {"时间": "", "医嘱内容": "", "执行情况": "", "执行人": ""}
  ],
  "nursing_measures": [
    {"时间": "", "护理措施": "", "患者反应": "", "记录人": ""}
  ],
  "handover_notes": "",
  "confidence": {}
}

## 提取规则
1. 时间格式统一为HH:MM。
2. 生命体征数值标准化为数字。
3. confidence字段：对每个记录项给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

PROMPT_RESEARCHER = """你是临床科研数据提取专家。请仔细识别该病历图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记、代码块标记或多余文字。JSON结构如下：

{
  "demographics": {
    "姓名": "", "性别": "", "年龄": null, "身高_cm": null, "体重_kg": null,
    "BMI": null, "婚姻状况": "", "职业": "", "民族": "",
    "吸烟史": "", "饮酒史": "", "过敏史": ""
  },
  "lab_tests": [
    {"项目名称": "", "英文缩写": "", "数值": null, "单位": "", "参考范围": "", "异常标注": ""}
  ],
  "treatment": {
    "入院时间": "", "出院时间": "", "住院天数": null, "科室": "",
    "主诊断": "", "主诊断ICD10编码": "", "其他诊断": [],
    "手术操作": "", "治疗方案": "", "出院医嘱": ""
  },
  "confidence": {
    "demographics": {}, "lab_tests": {}, "treatment": {}
  }
}

## 提取规则
1. 年龄提取纯数字（如"56岁"→56）；BMI如未给出尝试计算。
2. 实验室检查识别常见缩写：WBC、RBC、PLT、Hb、ALT、AST、Cr、BUN、GLU、TC、TG、HDL-C、LDL-C、UA、CRP、ESR、HbA1c、TSH、FT3、FT4、Na、K、Ca、Cl。异常标注"↑"/"↓"/"正常"。
3. 诊疗资料尽量识别ICD-10编码；治疗方案区分药物/手术/其他。
4. 置信度0-1，未识别到的字段为0。

## 医疗专业词汇参考
- 血常规：WBC、RBC、PLT、Hb、HCT、MCV、MCH、MCHC、RDW
- 肝功能：ALT、AST、GGT、ALP、TBIL、DBIL、TP、ALB
- 肾功能：Cr、BUN、UA、Cys-C、eGFR
- 血脂：TC、TG、HDL-C、LDL-C、ApoA1、ApoB
- 血糖：GLU、FPG、2hPG、HbA1c、OGTT
- 凝血：PT、APTT、TT、FIB、INR、D-Dimer
- 甲功：TSH、FT3、FT4、T3、T4
- 电解质：Na、K、Ca、Cl、Mg、P
- 炎症指标：CRP、PCT、ESR、IL-6
- 肿瘤标志物：AFP、CEA、CA125、CA199、CA153、PSA

只输出JSON，不要输出任何其他内容。"""

# 护士自定义模板的Prompt生成框架
NURSE_CUSTOM_PROMPT_TEMPLATE = """你是护理评估数据提取专家。请仔细识别该护理文档图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{{
  "patient_info": {{
    "姓名": "", "科室": "", "床号": "", "日期": ""
  }},
  "custom_fields": {{
    {field_schema}
  }},
  "confidence": {{}}
}}

## 提取规则
1. 识别图片中与以下字段相关的所有信息：{field_names}。
2. 数值型数据提取为数字，文本型数据提取为字符串。
{score_rule}
3. confidence字段：对每个已提取字段给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

# 医生自定义模板的Prompt生成框架
DOCTOR_CUSTOM_PROMPT_TEMPLATE = """你是临床病历数据提取专家。请仔细识别该医疗文档图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{{
  "patient_info": {{
    "姓名": "", "性别": "", "年龄": null, "科室": "", "床号": "", "住院号": ""
  }},
  "custom_fields": {{
    {field_schema}
  }},
  "confidence": {{}}
}}

## 提取规则
1. 识别图片中与以下字段相关的所有信息：{field_names}。
2. 年龄提取纯数字，诊断需尽量识别ICD-10编码。
3. 手写体请尽力识别，无法确认的字用?标记。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""

# 科研自定义模板的Prompt生成框架
RESEARCHER_CUSTOM_PROMPT_TEMPLATE = """你是临床科研数据提取专家。请仔细识别该医疗文档图片中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{{
  "demographics": {{
    "姓名": "", "性别": "", "年龄": null
  }},
  "custom_fields": {{
    {field_schema}
  }},
  "confidence": {{}}
}}

## 提取规则
1. 识别图片中与以下字段相关的所有信息：{field_names}。
2. 数值型数据提取为数字（如年龄、体重、血压数值等），文本型数据提取为字符串。
3. 日期格式统一为YYYY-MM-DD。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。

只输出JSON，不要输出任何其他内容。"""


# ========== 字段预览Prompt ==========

PROMPT_FIELD_PREVIEW = """你是医疗数据分析专家。请仔细分析以下医疗文档，识别所有可以提取的数据字段。

输出JSON格式:
{{
  "available_fields": [
    {{
      "field_name": "字段名称",
      "field_type": "text或number或date",
      "example_value": "从文档中提取的示例值",
      "confidence": 0.95,
      "category": "类别"
    }}
  ]
}}

要求:
1. 字段名称使用中文，简洁明确
2. 尽可能识别所有有意义的数据项（患者信息、检验指标、诊断、治疗、评估等）
3. category取值范围：基本信息、检验结果、诊疗记录、护理评估、其他
4. example_value必须是从文档中实际提取的真实值
5. confidence表示该字段在文档中的识别置信度(0-1)
6. 按category分组，同类别字段放在一起

只输出JSON，不要输出任何其他内容。"""

PROMPT_FIELD_PREVIEW_TEXT = """你是医疗数据分析专家。请仔细分析以下医疗文本，识别所有可以提取的数据字段。

输出JSON格式:
{{
  "available_fields": [
    {{
      "field_name": "字段名称",
      "field_type": "text或number或date",
      "example_value": "从文本中提取的示例值",
      "confidence": 0.95,
      "category": "类别"
    }}
  ]
}}

要求:
1. 字段名称使用中文，简洁明确
2. 尽可能识别所有有意义的数据项
3. category取值范围：基本信息、检验结果、诊疗记录、护理评估、其他
4. example_value必须是从文本中实际提取的真实值
5. confidence表示该字段的识别置信度(0-1)

文本内容:
{text_content}

只输出JSON，不要输出任何其他内容。"""

PROMPT_EXTRACT_FIELD_NAMES = """你是医疗模板设计专家，当前视角为{role_hint}。请仔细分析以下医疗相关文本，从中识别所有可以作为"数据提取模板字段名"的医学术语或概念。

注意：你要提取的是**字段名称**（如"血压"、"诊断"、"护理措施"），而不是具体的数据值。

输出JSON格式:
{{
  "fields": [
    {{
      "name": "字段名称",
      "category": "类别",
      "confidence": 0.95
    }}
  ]
}}

要求:
1. 字段名称使用中文，简洁明确（2-8个字为宜）
2. 只提取与医疗场景相关的字段，忽略无关词汇
3. category取值范围：基本信息、检验结果、诊疗记录、护理评估、科研数据、其他
4. confidence表示该词作为模板字段的合理程度(0-1)
5. 不要重复提取含义相同的字段
6. 根据{role_hint}的视角，优先识别该角色关注的字段
7. 不要将患者的具体姓名、具体数值等作为字段名

文本内容:
{text_content}

只输出JSON，不要输出任何其他内容。"""


# ========== 音频专用Prompt模板 ==========

PROMPT_AUDIO_DOCTOR = """你是临床医生数据提取专家。以下是医患对话的语音转录文本，请从中提取结构化病历信息。

## 输出格式(JSON):
{
  "patient_info": {"姓名": "", "性别": "", "年龄": null},
  "chief_complaint": "",
  "present_illness": "",
  "past_history": "",
  "physical_exam": {},
  "diagnosis": [{"诊断名称": "", "ICD10编码": ""}],
  "treatment_plan": {"药物治疗": "", "医嘱": "", "其他": ""},
  "conversation_notes": "",
  "confidence": {}
}

## 提取规则:
1. 从对话中识别患者自述的症状和病史
2. 提取医生口述的诊断和治疗建议
3. 主诉通常是患者开场描述的主要不适
4. 注意区分医生询问和患者回答
5. 如信息不完整或无法识别，对应字段留空字符串
6. confidence字段：对每个已提取字段给出0-1之间的置信度

只输出JSON，不要输出任何其他内容。"""

PROMPT_AUDIO_NURSE = """你是护理评估专家。以下是护理交班或患者访谈的语音转录文本，请提取护理相关信息。

## 输出格式(JSON):
{
  "patient_info": {"姓名": "", "床号": "", "科室": ""},
  "vital_signs_verbal": {},
  "nursing_observations": "",
  "patient_complaints": "",
  "nursing_actions": "",
  "handover_notes": "",
  "risk_alerts": "",
  "confidence": {}
}

## 提取规则:
1. 识别口述的生命体征数值（体温、血压、脉搏、呼吸、血氧等）
2. 提取护理观察内容（皮肤、伤口、活动能力、意识状态）
3. 记录患者主观感受和主诉
4. 提取交班时的重点提醒事项
5. 识别提及的护理风险（跌倒、压疮、管路等）
6. confidence字段：对每个已提取字段给出0-1之间的置信度

只输出JSON，不要输出任何其他内容。"""

PROMPT_AUDIO_RESEARCHER = """你是临床科研数据提取专家。以下是研究访谈或病历口述的语音转录文本，请提取科研相关数据。

## 输出格式(JSON):
{
  "demographics": {"姓名": "", "性别": "", "年龄": null, "职业": "", "教育程度": ""},
  "medical_history": "",
  "intervention_details": "",
  "outcome_measures": "",
  "patient_experience": "",
  "adherence_notes": "",
  "adverse_events": "",
  "research_notes": "",
  "confidence": {}
}

## 提取规则:
1. 提取人口学特征
2. 识别干预措施的描述
3. 提取患者自我报告的结局（症状改善、生活质量变化）
4. 注意提及的依从性和不良反应
5. 日期格式统一为YYYY-MM-DD
6. confidence字段：对每个已提取字段给出0-1之间的置信度

只输出JSON，不要输出任何其他内容。"""


# ========== 科室专属Prompt定义 ==========

# --- 心内科 ---
PROMPT_CARDIOLOGY_CLINICAL = """你是心内科数据提取专家。请识别该心内科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象，不要包含任何markdown标记或多余文字。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "既往史": "",
    "心电图": "",
    "心脏超声": "",
    "EF值": null,
    "BNP_NT_proBNP": null,
    "冠脉造影结果": "",
    "心律失常类型": "",
    "心功能分级_NYHA": "",
    "血压_收缩压_舒张压": "",
    "心率": null,
    "血脂_LDL_HDL_TG_TC": "",
    "肌钙蛋白": "",
    "用药方案": "",
    "PCI_CABG记录": "",
    "诊断": "",
    "治疗计划": ""
  },
  "confidence": {}
}

## 提取规则
1. EF值提取纯数字(百分比)。BNP/NT-proBNP提取数值及单位。
2. 冠脉造影需详细记录病变血管和狭窄程度。
3. 心功能分级按NYHA I-IV级标准。
4. 用药方案需区分抗血小板、他汀、ACEI/ARB、β受体阻滞剂等类别。
5. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_CARDIOLOGY_NURSING = """你是心内科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "心电监护": "",
    "生命体征_体温_脉搏_呼吸_血压": "",
    "胸痛评估": "",
    "出入量记录": "",
    "活动耐量评估": "",
    "抗凝药物管理": "",
    "跌倒风险评估": "",
    "心理状态": "",
    "饮食护理": "",
    "心脏康复指导": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 生命体征需提取完整(体温、脉搏、呼吸、血压)。
2. 胸痛评估包含部位、性质、持续时间、NRS评分。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_CARDIOLOGY_OTHER = """你是心血管科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "人口学特征": "",
    "LVEF": null,
    "BNP": null,
    "LDL_C": null,
    "支架类型": "",
    "再狭窄": "",
    "MACE事件": "",
    "随访日期": "",
    "用药情况": "",
    "治疗结局": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_CARDIOLOGY_AUDIO = """你是心内科医疗对话分析专家。请从该心内科相关的语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "心内科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "心电图提及": "",
    "EF值": null,
    "BNP": null,
    "用药方案": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 神经内科 ---
PROMPT_NEUROLOGY_CLINICAL = """你是神经内科数据提取专家。请识别该神经内科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "既往史": "",
    "GCS评分": null,
    "NIHSS评分": null,
    "肌力分级": "",
    "感觉障碍": "",
    "反射检查": "",
    "头颅CT_MRI": "",
    "脑电图": "",
    "腰穿结果": "",
    "病灶定位": "",
    "发病时间": "",
    "溶栓_取栓记录": "",
    "诊断": "",
    "治疗方案": ""
  },
  "confidence": {}
}

## 提取规则
1. GCS评分提取总分(3-15)和各项分值(E/V/M)。NIHSS提取总分(0-42)。
2. 肌力按0-5级标准，需注明部位(左上/左下/右上/右下)。
3. 发病时间精确到小时(用于评估溶栓时间窗)。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_NEUROLOGY_NURSING = """你是神经内科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "意识水平评估": "",
    "生命体征": "",
    "吞咽功能筛查": "",
    "跌倒风险评估": "",
    "压疮风险评估": "",
    "肢体活动度": "",
    "康复活动记录": "",
    "用药护理": "",
    "安全护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 意识水平需包含GCS评分或具体描述。
2. 吞咽筛查结果需明确通过/未通过及分级。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_NEUROLOGY_OTHER = """你是神经科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科"
  },
  "custom_fields": {
    "人口学特征": "",
    "mRS评分": null,
    "NIHSS评分": null,
    "发病至治疗时间": "",
    "梗死体积": "",
    "治疗方式": "",
    "复发情况": "",
    "功能预后": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_NEUROLOGY_AUDIO = """你是神经内科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "神经内科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "GCS评分": null,
    "NIHSS评分": null,
    "影像学提及": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 外科 ---
PROMPT_SURGERY_CLINICAL = """你是外科数据提取专家。请识别该外科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "既往史": "",
    "手术名称": "",
    "ASA分级": "",
    "麻醉方式": "",
    "切口类型": "",
    "术中出血量_ml": null,
    "术中输血": "",
    "手术时长_min": null,
    "引流管情况": "",
    "病理结果": "",
    "术后诊断": "",
    "出院医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 术中出血量和手术时长提取纯数字。
2. ASA分级按I-V级标准。切口类型分清洁/清洁-污染/污染/感染。
3. 引流管需记录类型、位置、引流量。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_SURGERY_NURSING = """你是外科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "术前评估": "",
    "生命体征": "",
    "切口护理": "",
    "引流管护理": "",
    "VTE预防": "",
    "疼痛评分_NRS": null,
    "术后活动": "",
    "饮食护理": "",
    "用药护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 切口护理需包含愈合分级(甲/乙/丙)和敷料情况。
2. 引流管护理需包含引流量、颜色、性状。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_SURGERY_OTHER = """你是外科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科"
  },
  "custom_fields": {
    "人口学特征": "",
    "手术方式": "",
    "手术时长_min": null,
    "出血量_ml": null,
    "并发症": "",
    "住院天数": null,
    "SSI发生": "",
    "再手术": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_SURGERY_AUDIO = """你是外科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "外科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "手术方案": "",
    "术后情况": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 儿科 ---
PROMPT_PEDIATRICS_CLINICAL = """你是儿科数据提取专家。请识别该儿科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "现病史": "",
    "出生体重_g": null,
    "胎龄_周": null,
    "体重_kg": null,
    "身高_cm": null,
    "头围_cm": null,
    "喂养方式": "",
    "疫苗接种史": "",
    "生长发育评估": "",
    "体格检查": "",
    "辅助检查": "",
    "诊断": "",
    "用药_剂量_kg": "",
    "治疗计划": ""
  },
  "confidence": {}
}

## 提取规则
1. 年龄需区分"岁"和"月龄"(如3岁2月、6月龄)。
2. 体重精确到0.1kg，身高精确到0.1cm。
3. 用药需注明按体重剂量(mg/kg)。
4. 疫苗接种记录需列出已接种和未接种疫苗。
5. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_PEDIATRICS_NURSING = """你是儿科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "体温调节": "",
    "喂养评估": "",
    "FLACC疼痛评分": null,
    "皮肤评估": "",
    "生命体征": "",
    "安全评估": "",
    "家长宣教": "",
    "排泄记录": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. FLACC疼痛评分适用于无法自述疼痛的儿童(0-10分)。
2. 喂养评估需包含奶量/辅食量、喂养耐受情况。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_PEDIATRICS_OTHER = """你是儿科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科"
  },
  "custom_fields": {
    "人口学特征": "",
    "身高百分位": "",
    "体重百分位": "",
    "BMI_Z评分": null,
    "发育里程碑": "",
    "营养评估": "",
    "疫苗完成率": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 百分位数需标注参考标准(WHO/CDC)。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_PEDIATRICS_AUDIO = """你是儿科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄_月龄": "", "科室": "儿科"
  },
  "custom_fields": {
    "主诉": "",
    "症状描述": "",
    "喂养情况": "",
    "发育情况": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 妇产科 ---
PROMPT_OBSTETRICS_CLINICAL = """你是妇产科数据提取专家。请识别该妇产科医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "主诉": "",
    "孕周": "",
    "孕次产次_GPAL": "",
    "末次月经": "",
    "胎心率": null,
    "宫高_cm": null,
    "腹围_cm": null,
    "B超_胎儿": "",
    "羊水指数": null,
    "血压": "",
    "尿蛋白": "",
    "分娩方式": "",
    "会阴状况": "",
    "Apgar评分": "",
    "新生儿体重_g": null,
    "诊断": "",
    "治疗计划": ""
  },
  "confidence": {}
}

## 提取规则
1. 孕周格式如"38+2周"。GPAL格式如"G2P1A0L1"。
2. Apgar评分需记录1分钟和5分钟值(如9-10-10)。
3. 胎心率正常范围110-160次/分。
4. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_OBSTETRICS_NURSING = """你是妇产科护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "产程监测": "",
    "宫缩评估": "",
    "胎心监护": "",
    "生命体征": "",
    "恶露观察": "",
    "母乳喂养评估": "",
    "产后出血评估": "",
    "会阴护理": "",
    "心理护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. 恶露需记录颜色、量、气味。
2. 产后出血评估需包含出血量估计。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_OBSTETRICS_OTHER = """你是妇产科科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科"
  },
  "custom_fields": {
    "人口学特征": "",
    "分娩孕周": "",
    "出生体重_g": null,
    "分娩方式": "",
    "产后出血量_ml": null,
    "并发症": "",
    "NICU入住": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_OBSTETRICS_AUDIO = """你是妇产科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "女", "年龄": null, "科室": "妇产科"
  },
  "custom_fields": {
    "主诉": "",
    "孕周": "",
    "症状描述": "",
    "胎心情况": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 急诊科 ---
PROMPT_EMERGENCY_CLINICAL = """你是急诊科数据提取专家。请识别该急诊医疗文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "分诊级别": "",
    "到达方式": "",
    "发病时间": "",
    "发病至就诊时间": "",
    "主诉": "",
    "MEWS_NEWS评分": null,
    "生命体征": "",
    "急救措施": "",
    "用药记录": "",
    "检查结果": "",
    "会诊记录": "",
    "绿色通道": "",
    "诊断": "",
    "转归": ""
  },
  "confidence": {}
}

## 提取规则
1. 分诊级别按I-IV级标准(I级濒危/II级危重/III级急症/IV级非急症)。
2. 发病至就诊时间精确到小时或分钟。
3. 转归需明确:留观/住院/出院/转院/死亡。
4. MEWS评分提取总分(0-14)。
5. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_EMERGENCY_NURSING = """你是急诊护理数据提取专家。请识别该护理文档中的所有文本内容，并严格按照以下JSON格式输出结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科", "床号": "", "住院号": ""
  },
  "custom_fields": {
    "ABCDE快速评估": "",
    "生命体征动态": "",
    "急救用药记录": "",
    "管路记录": "",
    "标本送检": "",
    "交接班核查": "",
    "疼痛评估": "",
    "安全护理": "",
    "护理问题": "",
    "护理措施": ""
  },
  "confidence": {}
}

## 提取规则
1. ABCDE评估包含气道/呼吸/循环/意识/暴露五项。
2. 管路记录需包含管路类型、置管时间、通畅情况。
3. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_EMERGENCY_OTHER = """你是急诊科研数据提取专家。请从该文档中提取结构化科研数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科"
  },
  "custom_fields": {
    "人口学特征": "",
    "就诊至处置时间_min": null,
    "分诊级别": "",
    "分诊准确率": "",
    "留观时长_h": null,
    "非计划重返": "",
    "转归": "",
    "随访日期": ""
  },
  "confidence": {}
}

## 提取规则
1. 数值型字段提取纯数字。日期统一YYYY-MM-DD格式。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

PROMPT_EMERGENCY_AUDIO = """你是急诊科医疗对话分析专家。请从该语音转录文本中提取结构化数据。

## 输出格式要求
请直接输出合法的JSON对象。JSON结构如下：

{
  "patient_info": {
    "姓名": "", "性别": "", "年龄": null, "科室": "急诊科"
  },
  "custom_fields": {
    "主诉": "",
    "发病时间": "",
    "症状描述": "",
    "急救措施": "",
    "诊断": "",
    "医嘱": ""
  },
  "confidence": {}
}

## 提取规则
1. 从对话中提取医学相关信息，忽略寒暄。
2. confidence字段：对每个已提取字段给出0-1之间的置信度。
只输出JSON，不要输出任何其他内容。"""

# --- 科室自动检测Prompt ---
PROMPT_DETECT_DEPARTMENT = """你是医疗科室分类专家。请分析以下医疗文本，判断它最可能属于哪个临床科室，并推荐合适的文档类型。

## 可选科室
1. cardiology (心内科) - 特征: 心电图、冠脉造影、EF值、BNP、NT-proBNP、心功能、PCI、房颤、心梗、心律失常、支架
2. neurology (神经内科) - 特征: GCS、NIHSS、偏瘫、脑梗、脑出血、癫痫、脑电图、腰穿、溶栓、卒中
3. surgery (外科) - 特征: 手术、切口、引流、麻醉、ASA、病理、术后、缝合、腹腔镜、清创
4. pediatrics (儿科) - 特征: 月龄、新生儿、疫苗、生长发育、喂养、辅食、百分位、胎龄
5. obstetrics (妇产科) - 特征: 孕周、产检、胎心、宫缩、分娩、剖宫产、恶露、Apgar、GPAL
6. emergency (急诊科) - 特征: 分诊、抢救、急救、120、绿色通道、MEWS、CPR、电除颤
7. general (通用) - 以上均不明确匹配时选择

## 文档类型
- clinical: 诊疗文档(病历、检查报告、诊断、治疗记录)
- nursing: 护理文档(护理评估、量表、护理记录)
- other: 其他(科研数据、访谈、综合报告)

## 输出格式
请直接输出合法的JSON对象：
{
  "department": "科室英文ID",
  "confidence": 0.0,
  "reasoning": "判断依据(30字以内)",
  "sub_type": "clinical或nursing或other"
}

只输出JSON，不要输出任何其他内容。

## 待分析文本：
"""

# 科室Prompt映射表（用于模板初始化和查找）
DEPARTMENT_PROMPTS = {
    'cardiology': {
        'clinical': PROMPT_CARDIOLOGY_CLINICAL,
        'nursing': PROMPT_CARDIOLOGY_NURSING,
        'other': PROMPT_CARDIOLOGY_OTHER,
        'audio': PROMPT_CARDIOLOGY_AUDIO,
    },
    'neurology': {
        'clinical': PROMPT_NEUROLOGY_CLINICAL,
        'nursing': PROMPT_NEUROLOGY_NURSING,
        'other': PROMPT_NEUROLOGY_OTHER,
        'audio': PROMPT_NEUROLOGY_AUDIO,
    },
    'surgery': {
        'clinical': PROMPT_SURGERY_CLINICAL,
        'nursing': PROMPT_SURGERY_NURSING,
        'other': PROMPT_SURGERY_OTHER,
        'audio': PROMPT_SURGERY_AUDIO,
    },
    'pediatrics': {
        'clinical': PROMPT_PEDIATRICS_CLINICAL,
        'nursing': PROMPT_PEDIATRICS_NURSING,
        'other': PROMPT_PEDIATRICS_OTHER,
        'audio': PROMPT_PEDIATRICS_AUDIO,
    },
    'obstetrics': {
        'clinical': PROMPT_OBSTETRICS_CLINICAL,
        'nursing': PROMPT_OBSTETRICS_NURSING,
        'other': PROMPT_OBSTETRICS_OTHER,
        'audio': PROMPT_OBSTETRICS_AUDIO,
    },
    'emergency': {
        'clinical': PROMPT_EMERGENCY_CLINICAL,
        'nursing': PROMPT_EMERGENCY_NURSING,
        'other': PROMPT_EMERGENCY_OTHER,
        'audio': PROMPT_EMERGENCY_AUDIO,
    },
}

# 科室子模板名称映射
DEPT_SUB_NAMES = {
    'clinical': '诊疗记录',
    'nursing': '护理评估',
    'other': '科研数据',
    'audio': '语音录入',
}

# 科室子模板display_layout映射
DEPT_SUB_LAYOUTS = {
    'clinical': 'table',
    'nursing': 'card',
    'other': 'table',
    'audio': 'table',
}


# ========== 数据库初始化 ==========
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db()
    c = conn.cursor()

    # 模板表
    c.execute('''CREATE TABLE IF NOT EXISTS extraction_templates (
        template_id TEXT PRIMARY KEY,
        role_id TEXT,
        template_name TEXT,
        template_type TEXT,
        ai_prompt TEXT,
        output_schema TEXT,
        display_layout TEXT,
        is_active INTEGER DEFAULT 1,
        create_time TEXT
    )''')

    # 记录表（新版）
    c.execute('''CREATE TABLE IF NOT EXISTS medical_records (
        id TEXT PRIMARY KEY,
        case_number TEXT UNIQUE,
        original_filename TEXT,
        role_id TEXT,
        template_id TEXT,
        extracted_data TEXT,
        confidence_data TEXT,
        demographics TEXT,
        lab_tests TEXT,
        treatment TEXT,
        confidence TEXT,
        raw_text TEXT,
        create_time TEXT
    )''')

    # 检查是否需要加新列（兼容旧数据库）
    existing_cols = {row[1] for row in c.execute("PRAGMA table_info(medical_records)").fetchall()}
    for col in ['role_id', 'template_id', 'extracted_data', 'confidence_data',
                'source_type', 'audio_transcript', 'qualitative_data',
                'module_type', 'text_source', 'analysis_type']:
        if col not in existing_cols:
            c.execute(f"ALTER TABLE medical_records ADD COLUMN {col} TEXT")

    # 迁移旧记录的module_type
    c.execute("UPDATE medical_records SET module_type='image_ocr' WHERE module_type IS NULL AND (source_type='image' OR source_type IS NULL)")
    c.execute("UPDATE medical_records SET module_type='voice_input' WHERE module_type IS NULL AND source_type='audio'")

    # 研究成果表
    c.execute('''CREATE TABLE IF NOT EXISTS research_results (
        result_id TEXT PRIMARY KEY,
        data_type TEXT,
        dept_id TEXT,
        source_record_id TEXT,
        title TEXT,
        summary TEXT,
        core_metrics TEXT,
        conclusion TEXT,
        notes TEXT,
        status TEXT DEFAULT '待复核',
        create_time TEXT,
        update_time TEXT
    )''')

    # 兼容旧表：如果 status 列不存在则添加
    try:
        c.execute("ALTER TABLE research_results ADD COLUMN status TEXT DEFAULT '待复核'")
    except Exception:
        pass

    conn.commit()
    conn.close()

    # 初始化内置模板
    _init_builtin_templates()
    # 插入演示数据
    _init_demo_data()


def _init_builtin_templates():
    """插入系统内置模板（如果尚未存在），使用INSERT OR IGNORE逐条插入"""
    conn = get_db()
    c = conn.cursor()

    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # ===== 数据迁移（幂等）：旧 role_id → 'general' =====
    for old_id in ('diagnosis', 'nursing', 'other', 'doctor', 'nurse', 'researcher'):
        c.execute("UPDATE extraction_templates SET role_id='general' WHERE role_id=?", (old_id,))
        c.execute("UPDATE medical_records SET role_id='general' WHERE role_id=?", (old_id,))

    # ===== 旧 12 个通用模板（role_id 统一改为 general）=====
    legacy_templates = [
        # 诊疗模板
        ('tpl_doctor_medical', 'general', '门诊/住院病历', 'fixed',
         PROMPT_DOCTOR_MEDICAL_RECORD, 'table', now),
        ('tpl_doctor_lab', 'general', '检查检验结果', 'fixed',
         PROMPT_DOCTOR_LAB_RESULTS, 'table', now),
        # 护理模板
        ('tpl_nurse_admission', 'general', '入院护理评估表', 'fixed',
         PROMPT_NURSE_ADMISSION, 'card', now),
        ('tpl_nurse_barthel', 'general', 'Barthel自理能力量表', 'fixed',
         PROMPT_NURSE_BARTHEL, 'scale', now),
        ('tpl_nurse_morse', 'general', 'Morse跌倒风险量表', 'fixed',
         PROMPT_NURSE_MORSE, 'scale', now),
        ('tpl_nurse_braden', 'general', 'Braden压疮风险量表', 'fixed',
         PROMPT_NURSE_BRADEN, 'scale', now),
        ('tpl_nurse_pain', 'general', 'NRS/VAS疼痛评估', 'fixed',
         PROMPT_NURSE_PAIN, 'card', now),
        ('tpl_nurse_record', 'general', '护理记录单', 'fixed',
         PROMPT_NURSE_RECORD, 'table', now),
        # 其他模板
        ('tpl_researcher_default', 'general', '综合科研数据提取', 'fixed',
         PROMPT_RESEARCHER, 'table', now),
        # 音频模板
        ('tpl_audio_doctor', 'general', '医患对话录音', 'fixed',
         PROMPT_AUDIO_DOCTOR, 'table', now),
        ('tpl_audio_nurse', 'general', '护理交班录音', 'fixed',
         PROMPT_AUDIO_NURSE, 'card', now),
        ('tpl_audio_researcher', 'general', '研究访谈录音', 'fixed',
         PROMPT_AUDIO_RESEARCHER, 'table', now),
    ]

    # ===== 新增 24 个科室专属模板 (6科室 × 4子模板) =====
    dept_templates = []
    for dept_id, sub_prompts in DEPARTMENT_PROMPTS.items():
        dept_name = DEPARTMENT_CONFIGS[dept_id]['name']
        for sub_type, prompt_text in sub_prompts.items():
            tpl_id = f"tpl_{dept_id}_{sub_type}"
            tpl_name = f"{dept_name} - {DEPT_SUB_NAMES[sub_type]}"
            layout = DEPT_SUB_LAYOUTS[sub_type]
            dept_templates.append(
                (tpl_id, dept_id, tpl_name, 'fixed', prompt_text, layout, now)
            )

    for t in legacy_templates + dept_templates:
        c.execute('''INSERT OR IGNORE INTO extraction_templates
            (template_id, role_id, template_name, template_type, ai_prompt, display_layout, create_time)
            VALUES (?, ?, ?, ?, ?, ?, ?)''', t)

    # 迁移旧数据：无 role_id 的记录归入 general
    c.execute('''UPDATE medical_records
        SET role_id='general', template_id='tpl_researcher_default'
        WHERE role_id IS NULL''')

    conn.commit()
    conn.close()


def _init_demo_data():
    """插入演示/测试数据（幂等：仅在 medical_records 为空时插入）"""
    conn = get_db()
    c = conn.cursor()

    count = c.execute("SELECT COUNT(*) FROM medical_records").fetchone()[0]
    if count > 0:
        conn.close()
        return

    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    demo_records = [
        {
            "id": "demo-rec-001",
            "case_number": "DEMO_20260316_A001",
            "original_filename": "心内科高血压病历.txt",
            "role_id": "cardiology",
            "template_id": "cardiology_diagnosis",
            "extracted_data": json.dumps({
                "患者姓名": "张**",
                "性别": "男",
                "年龄": "62岁",
                "主诉": "反复头晕、头痛3年，加重1周",
                "诊断": "原发性高血压3级（极高危）",
                "收缩压": "168mmHg",
                "舒张压": "102mmHg",
                "心率": "78bpm",
                "EF值": "58%",
                "BNP": "285pg/mL",
                "用药方案": "氨氯地平5mg qd + 缬沙坦80mg qd",
                "随访计划": "2周后复诊，监测血压日记"
            }, ensure_ascii=False),
            "raw_text": "患者张某某，男，62岁。因反复头晕、头痛3年，加重1周入院。既往高血压病史3年，最高血压180/110mmHg。入院查体：BP 168/102mmHg，HR 78bpm。心脏超声：EF 58%。BNP 285pg/mL。诊断：原发性高血压3级（极高危）。治疗方案：氨氯地平5mg qd + 缬沙坦80mg qd。",
            "source_type": "text",
            "module_type": "text_extract"
        },
        {
            "id": "demo-rec-002",
            "case_number": "DEMO_20260316_A002",
            "original_filename": "神经内科帕金森评估.txt",
            "role_id": "neurology",
            "template_id": "neurology_diagnosis",
            "extracted_data": json.dumps({
                "患者姓名": "李**",
                "性别": "男",
                "年龄": "71岁",
                "主诉": "右手静止性震颤2年，行走不稳6个月",
                "诊断": "帕金森病(H-Y 2.5期)",
                "UPDRS评分": "38分",
                "MMSE评分": "26分",
                "头颅MRI": "双侧基底节区少许缺血灶",
                "用药方案": "美多芭250mg tid + 普拉克索0.5mg tid",
                "alpha-synuclein": "阳性",
                "NfL": "32.5pg/mL"
            }, ensure_ascii=False),
            "raw_text": "患者李某某，男，71岁。右手静止性震颤2年，行走不稳6个月。查体：右侧肢体齿轮样肌强直，慌张步态。UPDRS评分38分，MMSE 26分。头颅MRI：双侧基底节区少许缺血灶。血清alpha-synuclein阳性，NfL 32.5pg/mL。诊断：帕金森病(H-Y 2.5期)。",
            "source_type": "text",
            "module_type": "text_extract"
        },
        {
            "id": "demo-rec-003",
            "case_number": "DEMO_20260316_A003",
            "original_filename": "儿科哮喘病历.txt",
            "role_id": "pediatrics",
            "template_id": "pediatrics_diagnosis",
            "extracted_data": json.dumps({
                "患者姓名": "王**",
                "性别": "女",
                "年龄": "6岁",
                "主诉": "反复喘息、咳嗽1年，加重3天",
                "诊断": "支气管哮喘（中度持续）",
                "肺功能FEV1": "72%预计值",
                "PEF变异率": "25%",
                "嗜酸性粒细胞": "6.2%",
                "IgE": "385IU/mL",
                "治疗方案": "布地奈德雾化吸入 1mg bid",
                "过敏原": "尘螨++，猫毛+"
            }, ensure_ascii=False),
            "raw_text": "患儿王某某，女，6岁。反复喘息、咳嗽1年，加重3天。肺功能：FEV1 72%预计值，PEF变异率25%。血常规：嗜酸性粒细胞6.2%。总IgE 385IU/mL。过敏原检测：尘螨++，猫毛+。诊断：支气管哮喘（中度持续）。治疗：布地奈德雾化吸入1mg bid。",
            "source_type": "text",
            "module_type": "text_extract"
        },
        {
            "id": "demo-rec-004",
            "case_number": "DEMO_20260316_A004",
            "original_filename": "急诊科胸痛分诊记录.txt",
            "role_id": "emergency",
            "template_id": "emergency_diagnosis",
            "extracted_data": json.dumps({
                "患者姓名": "赵**",
                "性别": "男",
                "年龄": "55岁",
                "主诉": "突发胸痛2小时",
                "分诊级别": "II级（紧急）",
                "MEWS评分": "4分",
                "心电图": "V1-V4 ST段抬高0.3-0.5mV",
                "肌钙蛋白I": "2.8ng/mL",
                "发病至就诊时间": "2小时",
                "诊断": "急性前壁ST段抬高型心肌梗死",
                "急救措施": "阿司匹林300mg+氯吡格雷300mg负荷，启动急诊PCI绿色通道",
                "转归": "急诊PCI成功，转入CCU"
            }, ensure_ascii=False),
            "raw_text": "患者赵某某，男，55岁。突发胸痛2小时，大汗淋漓。分诊II级。MEWS 4分。心电图：V1-V4 ST段抬高0.3-0.5mV。肌钙蛋白I 2.8ng/mL。诊断：急性前壁STEMI。负荷抗血小板，启动PCI绿色通道。Door-to-balloon 68min，PCI成功，转CCU。",
            "source_type": "text",
            "module_type": "text_extract"
        },
        {
            "id": "demo-rec-005",
            "case_number": "DEMO_20260316_A005",
            "original_filename": "外科术后感染评估.txt",
            "role_id": "surgery",
            "template_id": "surgery_diagnosis",
            "extracted_data": json.dumps({
                "患者姓名": "陈**",
                "性别": "女",
                "年龄": "48岁",
                "主诉": "右膝关节疼痛活动受限3年",
                "手术名称": "右侧全膝关节置换术",
                "ASA分级": "II级",
                "麻醉方式": "腰硬联合麻醉",
                "术中出血量": "280mL",
                "手术时间": "135分钟",
                "引流管": "负压引流管1根",
                "术后抗菌方案": "头孢唑林1g q8h x48h",
                "术后第3天体温": "37.2℃",
                "切口评估": "甲级愈合，无红肿渗出"
            }, ensure_ascii=False),
            "raw_text": "患者陈某某，女，48岁。右膝骨关节炎，行右侧TKA。ASA II级，腰硬联合麻醉。手术时间135min，出血280mL。术后头孢唑林1g q8h预防感染。术后第3天体温37.2℃，切口甲级愈合。",
            "source_type": "text",
            "module_type": "text_extract"
        }
    ]

    for rec in demo_records:
        c.execute('''INSERT OR IGNORE INTO medical_records
            (id, case_number, original_filename, role_id, template_id,
             extracted_data, raw_text, create_time, source_type, module_type)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (rec["id"], rec["case_number"], rec["original_filename"],
             rec["role_id"], rec["template_id"], rec["extracted_data"],
             rec["raw_text"], now, rec["source_type"], rec["module_type"]))

    conn.commit()
    conn.close()


# ========== 图片处理 ==========
def preprocess_image(image_path):
    try:
        img = Image.open(image_path)
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        gray = img.convert('L')
        gray = ImageOps.autocontrast(gray, cutoff=1)
        gray = gray.filter(ImageFilter.SHARPEN)
        enhancer = ImageEnhance.Contrast(gray)
        gray = enhancer.enhance(1.5)
        preprocessed_path = os.path.join(
            app.config['UPLOAD_FOLDER'],
            f"pre_{uuid.uuid4().hex[:8]}_{os.path.basename(image_path)}"
        )
        gray.save(preprocessed_path)
        return preprocessed_path
    except Exception:
        return image_path


def image_to_base64(image_path):
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode('utf-8')


def pdf_to_images(pdf_path):
    if not HAS_PYMUPDF:
        raise RuntimeError("PDF功能需要PyMuPDF库，请运行: pip install PyMuPDF")
    doc = fitz.open(pdf_path)
    image_paths = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        mat = fitz.Matrix(2.0, 2.0)
        pix = page.get_pixmap(matrix=mat)
        img_path = os.path.join(
            app.config['UPLOAD_FOLDER'],
            f"pdf_page_{uuid.uuid4().hex[:8]}_{page_num}.png"
        )
        pix.save(img_path)
        image_paths.append(img_path)
    doc.close()
    return image_paths


# ========== 本地OCR引擎 ==========

def local_ocr(image_path):
    """使用Tesseract对单张图片进行本地OCR识别"""
    if not HAS_TESSERACT:
        raise RuntimeError("pytesseract 未安装，请运行: pip install pytesseract，并确保系统已安装 Tesseract-OCR")
    preprocessed_path = preprocess_image(image_path)
    try:
        img = Image.open(preprocessed_path)
        text = pytesseract.image_to_string(img, lang='chi_sim+eng', config='--psm 6')
        return text.strip()
    finally:
        if preprocessed_path != image_path and os.path.exists(preprocessed_path):
            try:
                os.remove(preprocessed_path)
            except Exception:
                pass


def local_ocr_pdf(pdf_path):
    """对PDF的每一页进行OCR识别，合并全部文本"""
    image_paths = pdf_to_images(pdf_path)
    all_text = []
    for img_path in image_paths:
        try:
            page_text = local_ocr(img_path)
            if page_text:
                all_text.append(page_text)
        finally:
            try:
                if os.path.exists(img_path):
                    os.remove(img_path)
            except Exception:
                pass
    return '\n\n'.join(all_text)


def _extract_pdf_text(pdf_path):
    """使用PyMuPDF提取PDF中的嵌入文本（数字PDF直接提取，无需OCR）"""
    if not HAS_PYMUPDF:
        return ''
    try:
        doc = fitz.open(pdf_path)
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text and text.strip():
                text_parts.append(text.strip())
        doc.close()
        return '\n\n'.join(text_parts)
    except Exception:
        return ''


# ========== AI 识别核心 ==========

def extract_from_ocr_text(ocr_text, ai_prompt):
    """将本地OCR提取的文本发送给LLM进行结构化提取（发送前自动脱敏）"""
    # 脱敏：在发送远程LLM之前，对OCR文本进行敏感信息脱敏
    masked_text, _report = desensitize_text(ocr_text)
    combined_prompt = ai_prompt + "\n\n以下是通过OCR识别出的医疗文档文本，请按上述要求提取结构化信息：\n\n" + masked_text
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': combined_prompt}],
        temperature=0.1,
        max_tokens=4096
    )
    raw_text = response.choices[0].message.content
    parsed = parse_ai_response(raw_text)
    return parsed, raw_text


def extract_medical_data_multimodal(image_path, ai_prompt):
    """调用多模态模型直接识别图片（原始方式，作为OCR失败时的回退）
    注意：多模态模式下图片会发送至远程AI服务，返回的结构化数据将在本地进行脱敏。"""
    print("[PRIVACY] 多模态模式：图片将发送至AI服务进行识别")
    preprocessed_path = preprocess_image(image_path)
    b64_image = image_to_base64(preprocessed_path)

    ext = os.path.splitext(image_path)[1].lower()
    mime_map = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png', '.bmp': 'image/bmp',
        '.tiff': 'image/tiff'
    }
    mime_type = mime_map.get(ext, 'image/jpeg')

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{
            'role': 'user',
            'content': [
                {'type': 'text', 'text': ai_prompt},
                {
                    'type': 'image_url',
                    'image_url': {
                        'url': f"data:{mime_type};base64,{b64_image}"
                    }
                }
            ]
        }],
        temperature=0.1,
        max_tokens=4096
    )

    raw_text = response.choices[0].message.content

    try:
        if preprocessed_path != image_path and os.path.exists(preprocessed_path):
            os.remove(preprocessed_path)
    except Exception:
        pass

    parsed = parse_ai_response(raw_text)
    # 脱敏：对多模态返回的结构化数据进行脱敏
    parsed = desensitize_structured_data(parsed)
    return parsed, raw_text


def extract_medical_data(image_path, ai_prompt):
    """图片识别调度器：优先本地OCR + LLM结构化，失败则回退到多模态直接识别"""
    # 策略1: 尝试本地OCR
    if HAS_TESSERACT:
        try:
            ocr_text = local_ocr(image_path)
            if ocr_text and len(ocr_text) >= 10:
                print(f"[OCR] 本地OCR成功，提取文本长度: {len(ocr_text)}")
                parsed, raw_text = extract_from_ocr_text(ocr_text, ai_prompt)
                if 'error' not in parsed:
                    return parsed, raw_text
                print(f"[OCR] OCR文本结构化失败，回退到多模态识别")
            else:
                print(f"[OCR] 本地OCR文本过短({len(ocr_text) if ocr_text else 0}字符)，回退到多模态识别")
        except Exception as e:
            print(f"[OCR] 本地OCR失败: {e}，回退到多模态识别")

    # 策略2: 多模态模型直接识别
    return extract_medical_data_multimodal(image_path, ai_prompt)


def parse_ai_response(raw_text):
    text = re.sub(r'<think>.*?</think>', '', raw_text, flags=re.DOTALL)
    text = text.strip()
    if '```' in text:
        match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
        if match:
            text = match.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    start = text.find('{')
    end = text.rfind('}')
    if start != -1 and end != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            pass
    return {"error": "AI返回结果解析失败，请重试或检查图片质量", "raw_response": raw_text[:500]}


# ========== 语音识别与文本分析 ==========

def transcribe_audio(audio_path):
    """调用DashScope Paraformer识别本地音频文件"""
    if not HAS_DASHSCOPE:
        raise Exception("dashscope 库未安装，请运行 pip install dashscope")
    if not DASHSCOPE_API_KEY:
        raise Exception("环境变量 DASHSCOPE_API_KEY 未设置，无法使用语音识别")

    ext = os.path.splitext(audio_path)[1].lower()
    actual_path = audio_path
    converted = False

    # 不支持的格式转为wav
    if ext in ('.m4a', '.flac'):
        if not HAS_PYDUB:
            raise Exception("pydub 库未安装，无法转换 m4a/flac 格式，请运行 pip install pydub")
        audio_seg = AudioSegment.from_file(audio_path)
        actual_path = audio_path + '.wav'
        audio_seg.export(actual_path, format='wav')
        converted = True

    fmt_map = {'.wav': 'wav', '.mp3': 'mp3', '.aac': 'aac',
               '.amr': 'amr', '.opus': 'opus'}
    fmt = fmt_map.get(os.path.splitext(actual_path)[1].lower(), 'wav')

    try:
        recognition = Recognition(
            model='paraformer-realtime-v2',
            format=fmt,
            sample_rate=16000,
            language_hints=['zh', 'en']
        )
        result = recognition.call(actual_path)

        # 提取完整文本
        sentences = []
        if hasattr(result, 'get_sentence') and callable(result.get_sentence):
            sentences = result.get_sentence() or []
        full_text = ''.join([s.get('text', '') for s in sentences]) if sentences else ''

        if not full_text:
            # 尝试从output中获取
            if hasattr(result, 'output') and result.output:
                out = result.output
                if isinstance(out, dict) and 'text' in out:
                    full_text = out['text']
                elif isinstance(out, dict) and 'sentence' in out:
                    for s in out['sentence']:
                        full_text += s.get('text', '')

        if not full_text:
            raise Exception("语音识别未返回有效文本，请检查音频文件质量")

        return {
            'text': full_text,
            'sentences': sentences,
            'language': 'zh'
        }
    finally:
        if converted and os.path.exists(actual_path):
            try:
                os.remove(actual_path)
            except Exception:
                pass


def extract_from_transcript(transcript_text, ai_prompt):
    """用Qwen模型从转录文本中提取结构化数据（纯文本模式，发送前自动脱敏）"""
    # 脱敏：在发送远程LLM之前，对转录文本进行敏感信息脱敏
    masked_text, _report = desensitize_text(transcript_text)
    combined_prompt = ai_prompt + "\n\n以下是语音转录文本，请按上述要求提取结构化信息：\n\n" + masked_text
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': combined_prompt}],
        temperature=0.1,
        max_tokens=4096
    )
    raw_text = response.choices[0].message.content
    parsed = parse_ai_response(raw_text)
    return parsed, raw_text


# ========== 质性研究分析（科研角色专用） ==========

PROMPT_QUALITATIVE_ANALYSIS = """你是临床定性研究专家。请对以下医疗访谈转录文本进行定性分析。

输出JSON格式:
{{
  "themes": ["主题1", "主题2"],
  "keywords": ["关键词1", "关键词2"],
  "codes": [{{"code": "编码类别", "segments": ["相关文本片段1", "片段2"]}}],
  "sentiment": "积极/中性/消极",
  "summary": "2-3句话分析总结"
}}

分析要求:
1. 主题分析: 识别3-5个核心讨论主题
2. 关键词提取: 提取10-15个关键医学/情感词汇
3. 编码分类: 按类别(如症状描述、治疗态度、医患沟通、情感表达、生活影响)编码文本
4. 情感倾向: 判断整体情感

转录文本:
{transcript}

只输出JSON，不要输出任何其他内容。"""


def qualitative_analysis(transcript_text):
    """对转录文本进行定性研究分析（发送前自动脱敏）"""
    masked_text, _report = desensitize_text(transcript_text)
    prompt = PROMPT_QUALITATIVE_ANALYSIS.format(transcript=masked_text)
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': prompt}],
        temperature=0.3,
        max_tokens=3000
    )
    return parse_ai_response(response.choices[0].message.content)


# ========== 文本文件解析与预处理 ==========

def _parse_text_file(file_path):
    """解析txt/docx文件为纯文本"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.txt':
        for encoding in ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise Exception("无法识别文本文件编码，请使用UTF-8编码保存")

    elif ext == '.docx':
        try:
            import docx
        except ImportError:
            raise Exception("python-docx 库未安装，请运行 pip install python-docx")
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return '\n'.join(text_parts)

    elif ext == '.doc':
        raise Exception("不支持旧版.doc格式，请将文件另存为.docx后重新上传")

    else:
        raise Exception(f"不支持的文本格式: {ext}")


def _preprocess_text(text):
    """文本预处理：标准化格式"""
    if not text:
        return ''
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


# ========== 增强版质性研究分析 ==========

QUALITATIVE_TYPE_HINTS = {
    'interview': {
        'cn': '深度访谈',
        'hints': '患者体验、疾病认知、治疗态度、就医过程、心理感受',
        'coding': '症状描述、情感表达、医患沟通、治疗依从性、生活影响'
    },
    'focus_group': {
        'cn': '焦点小组',
        'hints': '群体共识、争议点、互动模式、观点演变、关键事件',
        'coding': '观点类别、互动类型、共识差异、关键事件、群体动态'
    },
    'observation': {
        'cn': '观察记录',
        'hints': '行为模式、环境因素、非语言信息、事件序列、场景特征',
        'coding': '行为类型、场景因素、时间特征、主体角色、环境条件'
    }
}

PROMPT_QUALITATIVE_ENHANCED = """你是质性研究方法论专家。请对以下{analysis_type_cn}材料严格按照四步质性分析法进行系统分析。

## 分析步骤

### 第一步：初始编码（Open Coding）
逐句或逐段阅读文本，识别有意义的概念、想法、行为模式，并为每个有意义的片段打上编码标签。
- 每个编码包含唯一编号(C01, C02...)、编码标签（简短概念名）、对应的原文片段、段落编号(P1, P2...)

### 第二步：主题聚类（Theme Clustering）
将相似的初始编码进行归纳合并，形成更高层次的上位主题类别。
- 识别3-5个主主题
- 每个主主题下有1-4个子主题
- 每个子主题关联具体的code_id列表
- 确保主题内部逻辑一致性和主题间差异性

### 第三步：典型原话保留（Representative Quotes）
为每个最终确定的主题选择2-3条最具代表性的原始引语。
- 必须是原文直接引用，不做改编
- 应能充分支撑该主题的核心观点

### 第四步：层级化输出（Hierarchical Structure）
最终输出严格按照"主题—子主题—编码—原话摘录"的层级结构。

## 输出JSON格式（严格遵守此结构）
{{
  "methodology_note": "本分析采用{analysis_type_cn}质性研究方法，遵循开放性编码→主题聚类→代表性引用→层级输出的四步分析流程",
  "analysis_type": "{analysis_type}",
  "step1_initial_coding": [
    {{
      "code_id": "C01",
      "code_label": "编码标签名",
      "original_text": "原文中的具体片段",
      "paragraph_ref": "P1"
    }}
  ],
  "step2_theme_clustering": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "sub_theme": "子主题名称",
          "codes": ["C01", "C03"],
          "description": "该子主题的简要描述"
        }}
      ]
    }}
  ],
  "step3_representative_quotes": [
    {{
      "theme": "主主题名称",
      "quotes": [
        "原话引用1",
        "原话引用2"
      ]
    }}
  ],
  "step4_hierarchical_output": [
    {{
      "theme": "主主题名称",
      "sub_themes": [
        {{
          "name": "子主题名称",
          "codes": [
            {{
              "label": "编码标签",
              "quotes": ["支撑该编码的原话"]
            }}
          ]
        }}
      ]
    }}
  ]
}}

## 分析要点
- 分析关注点：{analysis_hints}
- 编码参考类别：{coding_categories}
- 初始编码数量：8-20个（视文本长度而定）
- 主主题数量：3-5个
- 每个主题下的代表性引用：2-3条
- 所有引用必须来自原文，保持原始措辞

## 待分析文本：
{transcript}

只输出JSON，不要输出任何其他内容。"""


def qualitative_analysis_enhanced(transcript_text, analysis_type='interview'):
    """四步法质性分析：初始编码→主题聚类→代表性引用→层级输出（发送前自动脱敏）"""
    masked_text, _report = desensitize_text(transcript_text)
    type_info = QUALITATIVE_TYPE_HINTS.get(analysis_type, QUALITATIVE_TYPE_HINTS['interview'])
    prompt = PROMPT_QUALITATIVE_ENHANCED.format(
        analysis_type_cn=type_info['cn'],
        analysis_type=analysis_type,
        analysis_hints=type_info['hints'],
        coding_categories=type_info['coding'],
        transcript=masked_text
    )
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{'role': 'user', 'content': prompt}],
        temperature=0.3,
        max_tokens=4000
    )
    result = parse_ai_response(response.choices[0].message.content)
    # 基本校验：确保返回4步结构
    if 'error' not in result and 'step1_initial_coding' not in result:
        # 可能是旧格式返回，尝试包装
        if 'themes' in result:
            return result  # 返回旧格式，前端会兼容处理
    return result


# ========== 数据分析模块 ==========

def _extract_nested_field(data, field_path):
    """递归提取嵌套字段值，支持点分路径如 demographics.年龄"""
    parts = field_path.split('.', 1)
    if not isinstance(data, dict):
        return None
    val = data.get(parts[0])
    if len(parts) == 1:
        return val
    return _extract_nested_field(val, parts[1])


def _collect_field_paths(data, prefix=''):
    """递归收集JSON中所有叶子字段路径"""
    paths = []
    if isinstance(data, dict):
        for k, v in data.items():
            if k in ('confidence',):
                continue
            full = f"{prefix}.{k}" if prefix else k
            if isinstance(v, dict):
                paths.extend(_collect_field_paths(v, full))
            elif isinstance(v, list):
                if v and isinstance(v[0], dict):
                    paths.extend(_collect_field_paths(v[0], full + '[]'))
                else:
                    paths.append(full)
            else:
                paths.append(full)
    return paths


def _is_numeric(val):
    """判断值是否可转换为数值"""
    if val is None:
        return False
    try:
        float(val)
        return True
    except (ValueError, TypeError):
        return False


def analyze_structured_data(record_ids, fields, analysis_type='descriptive'):
    """对选中记录的结构化数据进行统计分析，返回统计量和ECharts配置"""
    conn = get_db()
    c = conn.cursor()
    placeholders = ','.join(['?'] * len(record_ids))
    c.execute(f'SELECT extracted_data, create_time FROM medical_records WHERE id IN ({placeholders})',
              record_ids)
    rows = c.fetchall()
    conn.close()

    # 收集每个字段的数值
    field_values = {f: [] for f in fields}
    time_series = {f: [] for f in fields}

    for row in rows:
        if not row['extracted_data']:
            continue
        data = json.loads(row['extracted_data'])
        create_time = row['create_time'] or ''
        for field in fields:
            val = _extract_nested_field(data, field)
            if _is_numeric(val):
                field_values[field].append(float(val))
                time_series[field].append({'time': create_time, 'value': float(val)})

    # 计算统计量
    stats = {}
    for field, values in field_values.items():
        if not values:
            stats[field] = {'count': 0, 'msg': '无有效数值'}
            continue
        arr = np.array(values)
        stats[field] = {
            'count': len(values),
            'mean': round(float(np.mean(arr)), 2),
            'median': round(float(np.median(arr)), 2),
            'std': round(float(np.std(arr)), 2),
            'min': round(float(np.min(arr)), 2),
            'max': round(float(np.max(arr)), 2),
            'q1': round(float(np.percentile(arr, 25)), 2),
            'q3': round(float(np.percentile(arr, 75)), 2)
        }

    # 生成ECharts配置
    chart_configs = []
    valid_fields = [f for f in fields if field_values[f]]

    if analysis_type == 'descriptive' and valid_fields:
        # 柱状图：各字段均值对比
        chart_configs.append({
            'title': {'text': '字段均值对比', 'left': 'center'},
            'tooltip': {'trigger': 'axis'},
            'xAxis': {'type': 'category', 'data': [f.split('.')[-1] for f in valid_fields],
                       'axisLabel': {'rotate': 30}},
            'yAxis': {'type': 'value', 'name': '均值'},
            'series': [{
                'data': [stats[f]['mean'] for f in valid_fields],
                'type': 'bar',
                'itemStyle': {'color': '#2563eb'},
                'label': {'show': True, 'position': 'top'}
            }],
            'grid': {'bottom': 80}
        })

        # 箱线图：分布概览
        if len(valid_fields) <= 8:
            boxplot_data = []
            for f in valid_fields:
                s = stats[f]
                boxplot_data.append([s['min'], s['q1'], s['median'], s['q3'], s['max']])
            chart_configs.append({
                'title': {'text': '数据分布（箱线图）', 'left': 'center'},
                'tooltip': {'trigger': 'item'},
                'xAxis': {'type': 'category', 'data': [f.split('.')[-1] for f in valid_fields],
                           'axisLabel': {'rotate': 30}},
                'yAxis': {'type': 'value'},
                'series': [{
                    'type': 'boxplot',
                    'data': boxplot_data,
                    'itemStyle': {'color': '#dbeafe', 'borderColor': '#2563eb'}
                }],
                'grid': {'bottom': 80}
            })

    elif analysis_type == 'trend' and valid_fields:
        # 折线图：按时间趋势
        series_list = []
        colors = ['#2563eb', '#059669', '#d97706', '#dc2626', '#7c3aed']
        for i, f in enumerate(valid_fields):
            sorted_ts = sorted(time_series[f], key=lambda x: x['time'])
            series_list.append({
                'name': f.split('.')[-1],
                'type': 'line',
                'data': [item['value'] for item in sorted_ts],
                'smooth': True,
                'itemStyle': {'color': colors[i % len(colors)]}
            })
        all_times = sorted(set(
            item['time'] for f in valid_fields for item in time_series[f]
        ))
        chart_configs.append({
            'title': {'text': '时间趋势分析', 'left': 'center'},
            'tooltip': {'trigger': 'axis'},
            'legend': {'data': [f.split('.')[-1] for f in valid_fields], 'bottom': 0},
            'xAxis': {'type': 'category', 'data': all_times, 'axisLabel': {'rotate': 45}},
            'yAxis': {'type': 'value'},
            'series': series_list,
            'grid': {'bottom': 80}
        })

    elif analysis_type == 'distribution' and valid_fields:
        # 直方图：频次分布（取第一个字段）
        f = valid_fields[0]
        values = field_values[f]
        n_bins = min(10, max(3, len(values) // 2))
        hist_counts, bin_edges = np.histogram(values, bins=n_bins)
        bin_labels = [f"{bin_edges[i]:.1f}-{bin_edges[i+1]:.1f}" for i in range(len(hist_counts))]
        chart_configs.append({
            'title': {'text': f'{f.split(".")[-1]} 频次分布', 'left': 'center'},
            'tooltip': {'trigger': 'axis'},
            'xAxis': {'type': 'category', 'data': bin_labels, 'axisLabel': {'rotate': 30}},
            'yAxis': {'type': 'value', 'name': '频次'},
            'series': [{
                'data': [int(c) for c in hist_counts],
                'type': 'bar',
                'itemStyle': {'color': '#059669'}
            }],
            'grid': {'bottom': 80}
        })

        # 饼图：分段占比
        if len(valid_fields) == 1:
            pie_data = [{'name': bin_labels[i], 'value': int(hist_counts[i])}
                        for i in range(len(hist_counts)) if hist_counts[i] > 0]
            chart_configs.append({
                'title': {'text': f'{f.split(".")[-1]} 分段占比', 'left': 'center'},
                'tooltip': {'trigger': 'item', 'formatter': '{b}: {c} ({d}%)'},
                'series': [{
                    'type': 'pie',
                    'radius': ['40%', '70%'],
                    'data': pie_data,
                    'label': {'formatter': '{b}\n{d}%'}
                }]
            })

    return {'statistics': stats, 'charts': chart_configs}


# ========== Excel 导出 ==========
def generate_excel(data_list):
    """多角色Excel导出，不同角色放不同Sheet"""
    from openpyxl.styles import Font, PatternFill
    red_fill = PatternFill(start_color='FFCCCC', end_color='FFCCCC', fill_type='solid')
    red_font = Font(color='CC0000', bold=True)

    # 按角色分组
    grouped = {}
    for item in data_list:
        role = item.get('role_id', 'general')
        role_name = DEPARTMENT_CONFIGS.get(role, {}).get('name', role)
        grouped.setdefault(role_name, []).append(item)

    excel_path = os.path.join(
        app.config['UPLOAD_FOLDER'],
        f"临床数据_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    )

    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        for sheet_name, items in grouped.items():
            rows = []
            low_conf_fields = set()

            for item in items:
                row = {
                    '病历编号': item.get('case_number', ''),
                    '模板': item.get('template_name', ''),
                    '录入时间': item.get('create_time', ''),
                    '数据来源': '录音' if item.get('source_type') == 'audio' else '图片',
                }
                data = item.get('extracted_data', {})
                conf = item.get('confidence_data', {})

                # 递归展平JSON数据为Excel列
                _flatten_to_row(data, row, '', conf, low_conf_fields)

                # 音频数据额外列
                if item.get('source_type') == 'audio' and item.get('audio_transcript'):
                    transcript = item['audio_transcript']
                    row['转录原文'] = transcript[:500] + ('...' if len(transcript) > 500 else '')
                qual = item.get('qualitative_data')
                if qual and isinstance(qual, dict):
                    row['主题分析'] = ', '.join(qual.get('themes', []))
                    row['关键词'] = ', '.join(qual.get('keywords', []))
                    row['情感倾向'] = qual.get('sentiment', '')

                rows.append(row)

            if not rows:
                continue

            df = pd.DataFrame(rows)
            safe_name = sheet_name[:31]  # Excel sheet名最长31字符
            df.to_excel(writer, index=False, sheet_name=safe_name)

            # 标红低置信度
            if low_conf_fields:
                ws = writer.sheets[safe_name]
                for col_idx, col_name in enumerate(df.columns, 1):
                    for field_name in low_conf_fields:
                        if field_name in col_name:
                            for row_idx in range(2, len(df) + 2):
                                cell = ws.cell(row=row_idx, column=col_idx)
                                cell.fill = red_fill
                                cell.font = red_font
                            ws.cell(row=1, column=col_idx).fill = red_fill
                            ws.cell(row=1, column=col_idx).font = red_font
                            break

    return excel_path


def _flatten_to_row(data, row, prefix, confidence, low_conf_fields):
    """递归展平嵌套JSON为Excel行的列"""
    if isinstance(data, dict):
        for k, v in data.items():
            if k == 'confidence':
                # 收集低置信度字段
                _collect_low_conf(v, low_conf_fields)
                continue
            full_key = f"{prefix}_{k}" if prefix else k
            if isinstance(v, dict):
                _flatten_to_row(v, row, full_key, confidence, low_conf_fields)
            elif isinstance(v, list):
                _flatten_list_to_row(v, row, full_key, confidence, low_conf_fields)
            else:
                row[full_key] = v
    elif isinstance(data, list):
        _flatten_list_to_row(data, row, prefix, confidence, low_conf_fields)


def _flatten_list_to_row(data_list, row, prefix, confidence, low_conf_fields):
    """展平列表数据"""
    for i, item in enumerate(data_list):
        if isinstance(item, dict):
            # 尝试用名称作为键
            name_key = item.get('项目名称') or item.get('英文缩写') or item.get('项目') or item.get('诊断名称') or str(i + 1)
            item_prefix = f"{prefix}_{name_key}" if prefix else name_key
            for k, v in item.items():
                if k in ('项目名称', '英文缩写', '项目', '诊断名称'):
                    continue
                if isinstance(v, (dict, list)):
                    continue
                col = f"{item_prefix}_{k}" if k != '数值' and k != '评分' else item_prefix
                row[col] = v
        elif isinstance(item, str):
            row[f"{prefix}_{i+1}"] = item


def _collect_low_conf(conf, low_conf_fields):
    """收集置信度<0.9的字段"""
    if isinstance(conf, dict):
        for k, v in conf.items():
            if isinstance(v, dict):
                _collect_low_conf(v, low_conf_fields)
            else:
                try:
                    if float(v) < 0.9:
                        low_conf_fields.add(k)
                except (ValueError, TypeError):
                    pass


# ========== 路由: 科室与模板 API ==========

@app.route('/api/roles', methods=['GET'])
def api_get_roles():
    """获取科室列表（兼容旧接口名 /api/roles）"""
    conn = get_db()
    c = conn.cursor()
    roles = []
    for dept_id, cfg in DEPARTMENT_CONFIGS.items():
        c.execute("SELECT COUNT(*) as cnt FROM extraction_templates WHERE role_id=? AND is_active=1",
                  (dept_id,))
        count = c.fetchone()['cnt']
        roles.append({
            'role_id': dept_id,
            'name': cfg['name'],
            'color': cfg['color'],
            'template_count': count
        })
    conn.close()
    return jsonify({"status": "success", "roles": roles})


@app.route('/api/departments', methods=['GET'])
def api_get_departments():
    """获取科室列表（新接口）"""
    conn = get_db()
    c = conn.cursor()
    departments = []
    for dept_id, cfg in DEPARTMENT_CONFIGS.items():
        c.execute("SELECT COUNT(*) as cnt FROM extraction_templates WHERE role_id=? AND is_active=1",
                  (dept_id,))
        count = c.fetchone()['cnt']
        departments.append({
            'dept_id': dept_id,
            'name': cfg['name'],
            'color': cfg['color'],
            'template_count': count
        })
    conn.close()
    return jsonify({"status": "success", "departments": departments})


@app.route('/api/templates/<role_id>', methods=['GET'])
def api_get_templates(role_id):
    """获取某科室下的模板列表"""
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT template_id, template_name, template_type, display_layout, ai_prompt, create_time
        FROM extraction_templates WHERE role_id=? AND is_active=1 ORDER BY template_type, create_time''',
              (role_id,))
    rows = c.fetchall()
    conn.close()
    templates = []
    for row in rows:
        fields = _extract_fields_from_prompt(row['ai_prompt']) if row['ai_prompt'] else []
        templates.append({
            'template_id': row['template_id'],
            'template_name': row['template_name'],
            'template_type': row['template_type'],
            'display_layout': row['display_layout'],
            'field_count': len(fields),
        })
    return jsonify({"status": "success", "templates": templates})


@app.route('/api/detect_department', methods=['POST'])
def api_detect_department():
    """AI自动检测医疗文本所属科室"""
    data = request.get_json()
    text = (data or {}).get('text', '').strip()
    if not text:
        return jsonify({"status": "error", "msg": "请提供待检测的医疗文本"})

    # 截取前2000字符送检以节省token
    sample = text[:2000]
    try:
        combined_prompt = PROMPT_DETECT_DEPARTMENT + "\n\n" + sample
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{'role': 'user', 'content': combined_prompt}],
            temperature=0.1,
            max_tokens=512
        )
        raw = response.choices[0].message.content
        parsed = parse_ai_response(raw)
        if isinstance(parsed, dict) and 'error' not in parsed:
            dept = parsed.get('department', 'general')
            # 校验科室ID有效性
            if dept not in DEPARTMENT_CONFIGS:
                dept = 'general'
            return jsonify({
                "status": "success",
                "department": dept,
                "confidence": parsed.get('confidence', 0),
                "reasoning": parsed.get('reasoning', ''),
                "sub_type": parsed.get('sub_type', 'clinical')
            })
        else:
            return jsonify({"status": "success", "department": "general",
                            "confidence": 0, "reasoning": "无法识别", "sub_type": "clinical"})
    except Exception as e:
        return jsonify({"status": "error", "msg": f"科室检测失败: {str(e)}",
                        "department": "general", "confidence": 0})


def _generate_template_prompt(role_id, fields, include_score=False, sub_type=None):
    """根据科室/角色和字段列表生成AI提取Prompt和display_layout。
    role_id 可以是新科室ID (cardiology等) 或旧角色ID (diagnosis/nursing/other)。
    sub_type 可选：'clinical'/'nursing'/'other'，用于科室模板。
    """
    field_schema_parts = []
    for f in fields:
        f = f.strip()
        if f:
            field_schema_parts.append(f'    "{f}": null')
    field_schema = ',\n'.join(field_schema_parts)
    field_names = '、'.join([f.strip() for f in fields if f.strip()])

    score_rule = ""
    if include_score:
        score_rule = "3. 如果字段是评分项，提取纯数字评分。如有总分，一并计算。\n"

    # 将旧角色ID映射为等效逻辑
    effective_role = LEGACY_ROLE_MAP.get(role_id, role_id)

    # 对于6大临床科室，根据 sub_type 选择模板风格
    if effective_role in DEPARTMENT_PROMPTS:
        dept_name = DEPARTMENT_CONFIGS.get(effective_role, {}).get('name', '临床')
        if sub_type == 'nursing':
            ai_prompt = NURSE_CUSTOM_PROMPT_TEMPLATE.format(
                field_schema=field_schema, field_names=field_names, score_rule=score_rule)
            display_layout = 'scale' if include_score else 'card'
        elif sub_type == 'other':
            ai_prompt = RESEARCHER_CUSTOM_PROMPT_TEMPLATE.format(
                field_schema=field_schema, field_names=field_names)
            display_layout = 'table'
        else:
            # clinical 或默认
            ai_prompt = DOCTOR_CUSTOM_PROMPT_TEMPLATE.format(
                field_schema=field_schema, field_names=field_names)
            display_layout = 'table'
    elif effective_role == 'general' or role_id == 'diagnosis':
        # 通用科室 / 旧诊疗角色 → 医生模板风格
        if sub_type == 'nursing' or role_id == 'nursing':
            ai_prompt = NURSE_CUSTOM_PROMPT_TEMPLATE.format(
                field_schema=field_schema, field_names=field_names, score_rule=score_rule)
            display_layout = 'scale' if include_score else 'card'
        elif sub_type == 'other' or role_id == 'other':
            ai_prompt = RESEARCHER_CUSTOM_PROMPT_TEMPLATE.format(
                field_schema=field_schema, field_names=field_names)
            display_layout = 'table'
        else:
            ai_prompt = DOCTOR_CUSTOM_PROMPT_TEMPLATE.format(
                field_schema=field_schema, field_names=field_names)
            display_layout = 'table'
    else:
        # 未知 role_id 兜底：护理风格
        ai_prompt = NURSE_CUSTOM_PROMPT_TEMPLATE.format(
            field_schema=field_schema, field_names=field_names, score_rule=score_rule)
        display_layout = 'scale' if include_score else 'card'

    return ai_prompt, display_layout


def _extract_fields_from_prompt(ai_prompt):
    """从ai_prompt中反向提取自定义字段列表"""
    fields = []
    matches = re.findall(r'"custom_fields"\s*:\s*\{([^}]+)\}', ai_prompt, re.DOTALL)
    if matches:
        field_matches = re.findall(r'"([^"]+)"\s*:\s*null', matches[0])
        fields = [f.strip() for f in field_matches if f.strip()]
    return fields


# 系统模板的中文字段名映射
TEMPLATE_FIELDS = {
    # ===== 通用（旧12模板）=====
    'tpl_doctor_medical': ['主诉', '现病史', '既往史', '个人史', '家族史', '过敏史', '体格检查', '专科检查', '辅助检查', '诊断', '诊疗计划', '处理意见'],
    'tpl_doctor_lab': ['项目名称', '结果', '参考值', '单位', '异常提示', '标本类型', '采集时间', '报告时间'],
    'tpl_nurse_admission': ['一般资料', '过敏史', '既往史', '用药史', '生命体征', '意识状态', '皮肤黏膜', '营养状况', '排泄', '活动能力', '跌倒风险', '压疮风险', '疼痛评分', '吞咽功能', '心理状态', '睡眠', '饮食', '专科情况', '护理问题', '护理措施'],
    'tpl_nurse_barthel': ['进食', '洗澡', '修饰', '穿衣', '控制大便', '控制小便', '如厕', '床椅转移', '平地行走', '上下楼梯'],
    'tpl_nurse_morse': ['跌倒史', '继发诊断', '步行辅助', '静脉输液/肝素锁', '步态', '认知状态'],
    'tpl_nurse_braden': ['感知能力', '潮湿程度', '活动能力', '移动能力', '营养摄取', '摩擦力和剪切力'],
    'tpl_nurse_pain': ['疼痛部位', '疼痛性质', '疼痛强度', '诱发因素', '缓解因素', '伴随症状', '疼痛持续时间'],
    'tpl_nurse_record': ['生命体征', '意识状态', '皮肤完整性', '跌倒风险', '压疮风险', '护理措施'],
    'tpl_researcher_default': ['人口学特征', '实验室检查', '主要终点事件', '随访日期', '血压', '血脂', '用药情况', '治疗结局'],
    'tpl_audio_doctor': ['主诉', '现病史', '诊断', '治疗方案'],
    'tpl_audio_nurse': ['生命体征', '护理观察', '风险提醒'],
    'tpl_audio_researcher': ['人口学特征', '病史', '干预措施', '结局指标'],
    # ===== 心内科 =====
    'tpl_cardiology_clinical': ['主诉', '现病史', '既往史', 'EF值', 'BNP_NT_proBNP', '冠脉造影结果', '心律失常类型', '心功能分级_NYHA', '血压_收缩压_舒张压', '心率', '血脂_LDL_HDL_TG_TC', '肌钙蛋白', '用药方案', 'PCI_CABG记录', '诊断', '治疗计划'],
    'tpl_cardiology_nursing': ['心电监护', '生命体征_体温_脉搏_呼吸_血压', '胸痛评估', '出入量记录', '活动耐量评估', '抗凝药物管理', '跌倒风险评估', '心理状态', '饮食护理', '心脏康复指导', '护理问题', '护理措施'],
    'tpl_cardiology_other': ['人口学特征', 'LVEF', 'BNP', 'LDL_C', '支架类型', '再狭窄', 'MACE事件', '随访日期', '用药情况', '治疗结局'],
    'tpl_cardiology_audio': ['主诉', '症状描述', '心电图提及', 'EF值', 'BNP', '用药方案', '诊断', '医嘱'],
    # ===== 神经内科 =====
    'tpl_neurology_clinical': ['主诉', '现病史', '既往史', 'GCS评分', 'NIHSS评分', '肌力分级', '感觉障碍', '反射检查', '头颅CT_MRI', '脑电图', '腰穿结果', '病灶定位', '发病时间', '溶栓_取栓记录', '诊断', '治疗方案'],
    'tpl_neurology_nursing': ['意识状态_GCS', '瞳孔变化', '肌力评估', '吞咽功能筛查', '跌倒风险', '深静脉血栓预防', '语言功能评估', '康复训练', '颅内压监测', '护理问题', '护理措施'],
    'tpl_neurology_other': ['人口学特征', '卒中类型', 'NIHSS评分', '发病至治疗时间', 'mRS评分', '影像学结果', '再灌注治疗', '并发症', '随访日期', '治疗结局'],
    'tpl_neurology_audio': ['主诉', '症状描述', '意识状态', '肢体活动', '影像学提及', '诊断', '医嘱'],
    # ===== 外科 =====
    'tpl_surgery_clinical': ['主诉', '现病史', '既往史', '手术名称', 'ASA分级', '麻醉方式', '术中出血量', '手术时长', '引流管', '术后并发症', '病理结果', '切口情况', '诊断', '手术记录'],
    'tpl_surgery_nursing': ['术前准备', '手术核查', '术后生命体征', '引流管护理', '切口换药', '疼痛评估', 'DVT预防', '早期活动', '营养支持', '术后并发症观察', '护理问题', '护理措施'],
    'tpl_surgery_other': ['人口学特征', '手术方式', '手术时长', '术中出血量', '并发症', '住院天数', '切口愈合等级', '病理分期', '随访日期', '治疗结局'],
    'tpl_surgery_audio': ['主诉', '症状描述', '手术方案', '麻醉方式', '术后情况', '诊断', '医嘱'],
    # ===== 儿科 =====
    'tpl_pediatrics_clinical': ['主诉', '现病史', '既往史', '出生体重', '胎龄', '体重_身高百分位', '疫苗接种', '喂养方式', '发育评估', '过敏史', '家族史', '专科检查', '诊断', '治疗方案'],
    'tpl_pediatrics_nursing': ['体温管理', '喂养评估', '体重监测', '黄疸评估', '疼痛评估_FLACC', '用药核查_体重剂量', '家长健康教育', '隔离防护', '护理问题', '护理措施'],
    'tpl_pediatrics_other': ['人口学特征', '胎龄', '出生体重', '生长发育指标', '疫苗接种', '过敏', '主要诊断', '治疗方案', '随访日期', '治疗结局'],
    'tpl_pediatrics_audio': ['主诉', '症状描述', '体温', '喂养情况', '发育情况', '诊断', '医嘱'],
    # ===== 妇产科 =====
    'tpl_obstetrics_clinical': ['主诉', '现病史', '既往史', '孕周', 'GPAL', '胎心率', '宫高_腹围', '羊水指数', '胎位', '妊娠并发症', '分娩方式', 'Apgar评分', '产后出血量', '新生儿体重', '诊断', '治疗方案'],
    'tpl_obstetrics_nursing': ['产前监护', '胎心监测', '宫缩评估', '产后出血观察', '母乳喂养指导', '切口_会阴护理', '子宫复旧', '新生儿护理', '心理支持', '护理问题', '护理措施'],
    'tpl_obstetrics_other': ['人口学特征', '孕周', '分娩方式', '妊娠并发症', '新生儿结局', 'Apgar评分', '产后出血量', '住院天数', '随访日期', '治疗结局'],
    'tpl_obstetrics_audio': ['主诉', '症状描述', '孕周', '胎动情况', '产检结果', '诊断', '医嘱'],
    # ===== 急诊科 =====
    'tpl_emergency_clinical': ['主诉', '现病史', '既往史', '分诊级别', 'MEWS_NEWS评分', '发病至就诊时间', '生命体征', '意识状态_GCS', '急救措施', '检查结果', '会诊情况', '处置结果', '转归', '诊断'],
    'tpl_emergency_nursing': ['分诊评估', '生命体征', '意识状态', '疼痛评估', '急救配合', '静脉通路', '用药记录', '标本采集', '转运交接', '护理问题', '护理措施'],
    'tpl_emergency_other': ['人口学特征', '就诊时间', '分诊级别', '主要诊断', '急救措施', '检查项目', '抢救结局', '滞留时间', '随访日期', '治疗结局'],
    'tpl_emergency_audio': ['主诉', '症状描述', '发病时间', '急救措施', '意识状态', '诊断', '医嘱'],
}


@app.route('/api/templates/<template_id>/detail', methods=['GET'])
def api_get_template_detail(template_id):
    """获取模板完整信息用于编辑"""
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT template_id, role_id, template_name, template_type,
        ai_prompt, display_layout, create_time
        FROM extraction_templates WHERE template_id=?''', (template_id,))
    row = c.fetchone()
    conn.close()
    if not row:
        return jsonify({"status": "error", "msg": "模板不存在"})

    # For system templates, try extracting from updated prompt first (user may have edited fields),
    # then fall back to TEMPLATE_FIELDS defaults
    fields = []
    if row['ai_prompt']:
        fields = _extract_fields_from_prompt(row['ai_prompt'])
    if not fields:
        fields = TEMPLATE_FIELDS.get(row['template_id'], [])
    include_score = row['display_layout'] == 'scale'

    return jsonify({
        "status": "success",
        "template": {
            "template_id": row['template_id'],
            "role_id": row['role_id'],
            "template_name": row['template_name'],
            "template_type": row['template_type'],
            "display_layout": row['display_layout'],
            "fields": fields,
            "include_score": include_score,
            "create_time": row['create_time']
        }
    })


@app.route('/api/templates', methods=['POST'])
def api_create_template():
    """创建自定义模板"""
    data = request.get_json()
    if not data:
        return jsonify({"status": "error", "msg": "无数据"})

    role_id = data.get('role_id', 'general')
    template_name = data.get('template_name', '').strip()
    fields = data.get('fields', [])
    include_score = data.get('include_score', False)
    sub_type = data.get('sub_type', None)

    if not template_name or not fields:
        return jsonify({"status": "error", "msg": "请填写模板名称和提取字段"})

    ai_prompt, display_layout = _generate_template_prompt(role_id, fields, include_score, sub_type=sub_type)

    template_id = f"tpl_custom_{uuid.uuid4().hex[:8]}"
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT INTO extraction_templates
        (template_id, role_id, template_name, template_type, ai_prompt, display_layout, is_active, create_time)
        VALUES (?, ?, ?, 'custom', ?, ?, 1, ?)''',
              (template_id, role_id, template_name, ai_prompt, display_layout, now))
    conn.commit()
    conn.close()

    return jsonify({"status": "success", "template_id": template_id, "msg": "模板创建成功"})


@app.route('/api/templates/<template_id>', methods=['DELETE'])
def api_delete_template(template_id):
    """删除自定义模板"""
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT template_type FROM extraction_templates WHERE template_id=?", (template_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"status": "error", "msg": "模板不存在"})
    if row['template_type'] == 'fixed':
        conn.close()
        return jsonify({"status": "error", "msg": "系统内置模板不可删除"})

    c.execute("DELETE FROM extraction_templates WHERE template_id=?", (template_id,))
    conn.commit()
    conn.close()
    return jsonify({"status": "success", "msg": "模板已删除"})


@app.route('/api/templates/<template_id>', methods=['PUT'])
def api_update_template(template_id):
    """编辑自定义模板"""
    data = request.get_json()
    if not data:
        return jsonify({"status": "error", "msg": "无数据"})

    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT template_type, role_id FROM extraction_templates WHERE template_id=?", (template_id,))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({"status": "error", "msg": "模板不存在"})

    role_id = row['role_id']
    template_name = data.get('template_name', '').strip()
    fields = data.get('fields', [])
    include_score = data.get('include_score', False)
    sub_type = data.get('sub_type', None)

    # For system templates, keep original name if not provided
    if not template_name and row['template_type'] == 'fixed':
        c.execute("SELECT template_name FROM extraction_templates WHERE template_id=?", (template_id,))
        name_row = c.fetchone()
        template_name = name_row['template_name'] if name_row else ''

    if not template_name or not fields:
        conn.close()
        return jsonify({"status": "error", "msg": "请填写模板名称和提取字段"})

    ai_prompt, display_layout = _generate_template_prompt(role_id, fields, include_score, sub_type=sub_type)

    c.execute('''UPDATE extraction_templates
        SET template_name=?, ai_prompt=?, display_layout=?
        WHERE template_id=?''',
              (template_name, ai_prompt, display_layout, template_id))
    conn.commit()
    conn.close()

    # Update in-memory fields cache
    TEMPLATE_FIELDS[template_id] = fields

    return jsonify({"status": "success", "msg": "模板已更新"})


# ========== 字段预览与自定义提取 ==========

@app.route('/api/extract_fields_from_text', methods=['POST'])
def api_extract_fields_from_text():
    """从用户输入的描述性文本中智能提取可用作模板字段的名称"""
    try:
        data = request.get_json()
        text = (data.get('text', '') or '').strip()
        role_id = data.get('role_id', 'general')

        if len(text) < 5:
            return jsonify({"status": "error", "msg": "请输入更多文本内容（至少5个字符）"})

        if role_id not in ('diagnosis', 'nursing', 'other'):
            role_id = 'other'

        role_hints = {
            'diagnosis': '诊疗数据提取',
            'nursing': '护理评估数据提取',
            'other': '综合数据提取'
        }
        role_hint = role_hints[role_id]

        # 脱敏处理
        masked_text, _report = desensitize_text(text)

        # 构建 prompt 并调用 LLM
        prompt = PROMPT_EXTRACT_FIELD_NAMES.format(
            role_hint=role_hint,
            text_content=masked_text
        )

        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{'role': 'user', 'content': prompt}],
            temperature=0.1,
            max_tokens=2048
        )
        raw_text = response.choices[0].message.content
        parsed = parse_ai_response(raw_text)

        # 提取字段列表
        fields = parsed.get('fields', [])
        if not isinstance(fields, list):
            fields = []

        # 过滤低置信度字段并去重
        seen = set()
        filtered = []
        category_order = {'基本信息': 0, '检验结果': 1, '诊疗记录': 2, '护理评估': 3, '科研数据': 4, '其他': 5}
        for f in fields:
            if not isinstance(f, dict):
                continue
            name = (f.get('name', '') or '').strip()
            confidence = f.get('confidence', 0)
            if not name or name in seen:
                continue
            if isinstance(confidence, (int, float)) and confidence < 0.5:
                continue
            seen.add(name)
            filtered.append({
                'name': name,
                'category': f.get('category', '其他'),
                'confidence': round(confidence, 2) if isinstance(confidence, (int, float)) else 0.8
            })

        # 按类别排序
        filtered.sort(key=lambda x: category_order.get(x['category'], 5))

        return jsonify({"status": "success", "fields": filtered})

    except Exception as e:
        return jsonify({"status": "error", "msg": f"分析失败: {str(e)}"})


@app.route('/api/preview_fields', methods=['POST'])
def api_preview_fields():
    """文档字段预览 - 分析文档并返回可提取的字段列表"""
    text_content = request.form.get('text_content', '').strip()
    uploaded_files = request.files.getlist('files')

    raw_data = None
    fields = []

    try:
        if text_content:
            # 文本模式
            prompt = PROMPT_FIELD_PREVIEW_TEXT.format(text_content=text_content)
            response = client.chat.completions.create(
                model=MODEL_NAME,
                messages=[{'role': 'user', 'content': prompt}],
                temperature=0.1,
                max_tokens=4096
            )
            raw_data = parse_ai_response(response.choices[0].message.content)

        elif uploaded_files and uploaded_files[0].filename:
            file = uploaded_files[0]
            file_ext = os.path.splitext(file.filename)[1].lower()
            temp_name = f"{uuid.uuid4().hex}{file_ext}"
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_name)
            file.save(file_path)

            try:
                if is_audio_file(file.filename):
                    transcript_result = transcribe_audio(file_path)
                    transcript_text = transcript_result['text']
                    prompt = PROMPT_FIELD_PREVIEW_TEXT.format(text_content=transcript_text)
                    response = client.chat.completions.create(
                        model=MODEL_NAME,
                        messages=[{'role': 'user', 'content': prompt}],
                        temperature=0.1,
                        max_tokens=4096
                    )
                    raw_data = parse_ai_response(response.choices[0].message.content)
                elif is_text_file(file.filename):
                    raw_file_text = _parse_text_file(file_path)
                    processed_text = _preprocess_text(raw_file_text)
                    prompt = PROMPT_FIELD_PREVIEW_TEXT.format(text_content=processed_text)
                    response = client.chat.completions.create(
                        model=MODEL_NAME,
                        messages=[{'role': 'user', 'content': prompt}],
                        temperature=0.1,
                        max_tokens=4096
                    )
                    raw_data = parse_ai_response(response.choices[0].message.content)
                else:
                    # 图片/PDF模式
                    if file_ext == '.pdf':
                        image_paths = pdf_to_images(file_path)
                        if image_paths:
                            raw_data, _ = extract_medical_data(image_paths[0], PROMPT_FIELD_PREVIEW)
                            for ip in image_paths:
                                try:
                                    os.remove(ip)
                                except Exception:
                                    pass
                    else:
                        raw_data, _ = extract_medical_data(file_path, PROMPT_FIELD_PREVIEW)
            finally:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                except Exception:
                    pass
        else:
            return jsonify({"status": "error", "msg": "请提供文件或文本内容"})

        if not raw_data or 'error' in raw_data:
            return jsonify({"status": "error", "msg": raw_data.get('error', '分析失败') if raw_data else '分析失败'})

        fields = raw_data.get('available_fields', [])
        # 过滤低置信度字段
        fields = [f for f in fields if f.get('confidence', 0) >= 0.5]
        # 按category排序
        category_order = {'基本信息': 0, '检验结果': 1, '诊疗记录': 2, '护理评估': 3, '其他': 4}
        fields.sort(key=lambda x: category_order.get(x.get('category', '其他'), 4))

        return jsonify({
            "status": "success",
            "fields": fields,
            "raw_data": raw_data
        })
    except Exception as e:
        return jsonify({"status": "error", "msg": f"字段预览失败: {str(e)}"})


@app.route('/api/extract_selected', methods=['POST'])
def api_extract_selected_fields():
    """根据用户选择的字段执行提取"""
    data = request.get_json()
    if not data:
        return jsonify({"status": "error", "msg": "无数据"})

    selected_fields = data.get('selected_fields', [])
    role_id = data.get('role_id', 'general')
    sub_type = data.get('sub_type', None)
    cached_raw_data = data.get('raw_data')
    text_content = data.get('text_content', '').strip()

    if not selected_fields:
        return jsonify({"status": "error", "msg": "请至少选择一个字段"})

    results = []
    errors = []

    try:
        extracted_data = {}

        if cached_raw_data:
            # 从缓存中筛选字段
            all_fields = cached_raw_data.get('available_fields', [])
            for f in all_fields:
                if f.get('field_name') in selected_fields:
                    extracted_data[f['field_name']] = f.get('example_value')
        elif text_content:
            # 脱敏：在发送远程LLM之前，对文本进行敏感信息脱敏
            masked_content, _privacy_report = desensitize_text(text_content)
            # 动态生成prompt提取
            ai_prompt, _ = _generate_template_prompt(role_id, selected_fields, sub_type=sub_type)
            parsed, raw_text = extract_from_transcript(masked_content, ai_prompt)
            if 'error' not in parsed:
                extracted_data = parsed.get('custom_fields', parsed)
            else:
                return jsonify({"status": "error", "msg": parsed.get('error', '提取失败')})
        else:
            return jsonify({"status": "error", "msg": "缺少数据来源"})

        # 存储到数据库
        case_number = f"DYN_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
        record_id = str(uuid.uuid4())
        create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        conn = get_db()
        c = conn.cursor()
        c.execute('''INSERT INTO medical_records
            (id, case_number, original_filename, role_id, template_id,
             extracted_data, confidence_data, raw_text, create_time,
             source_type, module_type)
            VALUES (?, ?, ?, ?, 'dynamic_extract', ?, NULL, NULL, ?, 'text', 'dynamic_extract')''',
            (record_id, case_number, '自定义字段提取', role_id,
             json.dumps(extracted_data, ensure_ascii=False), create_time))
        conn.commit()
        conn.close()

        results.append({
            "id": record_id,
            "case_number": case_number,
            "filename": "自定义字段提取",
            "role_id": role_id,
            "template_name": "自定义字段",
            "display_layout": "table",
            "source_type": "text",
            "module_type": "dynamic_extract",
            "data": extracted_data,
            "create_time": create_time
        })
    except Exception as e:
        errors.append(f"提取失败: {str(e)}")

    return jsonify({
        "status": "success" if results else "error",
        "results": results,
        "errors": errors,
        "msg": f"成功提取 {len(results)} 份" if results else "提取失败"
    })


# ========== 路由: 核心功能 ==========

@app.route('/')
@app.route('/index')
def home():
    init_db()
    return render_template('home.html')

@app.route('/data-extraction')
def data_extraction():
    init_db()
    return render_template('data_extraction.html')

@app.route('/research-results')
def research_results_page():
    init_db()
    return render_template('research_results.html')


@app.route('/upload', methods=['POST'])
def upload_and_recognize():
    """上传文件并AI识别（支持角色/模板选择，支持图片和音频）"""
    if 'files' not in request.files:
        return jsonify({"status": "error", "msg": "未选择文件"})

    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        return jsonify({"status": "error", "msg": "未选择有效文件"})

    role_id = request.form.get('role_id', 'general')
    template_id = request.form.get('template_id', 'tpl_researcher_default')
    module_type = request.form.get('module_type', '')  # 由前端指定

    # 查询模板Prompt
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT ai_prompt, template_name, display_layout FROM extraction_templates WHERE template_id=?",
              (template_id,))
    tpl_row = c.fetchone()
    conn.close()

    if not tpl_row:
        return jsonify({"status": "error", "msg": "模板不存在"})

    ai_prompt = tpl_row['ai_prompt']
    template_name = tpl_row['template_name']
    display_layout = tpl_row['display_layout']

    results = []
    errors = []

    for file in files:
        is_audio = is_audio_file(file.filename)
        is_image = allowed_file(file.filename)

        if not is_audio and not is_image:
            errors.append(f"不支持的格式: {file.filename}")
            continue

        file_ext = os.path.splitext(file.filename)[1].lower()
        temp_name = f"{uuid.uuid4().hex}{file_ext}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_name)
        file.save(file_path)

        try:
            if is_audio:
                # ========== 音频处理流程 ==========
                result_data = _process_audio_file(
                    file_path, file.filename, role_id, template_id,
                    ai_prompt, template_name, display_layout)
                if result_data.get('error'):
                    errors.append(f"{file.filename}: {result_data['error']}")
                else:
                    results.append(result_data)
            else:
                # ========== 图片/PDF处理流程（本地OCR + 远程AI增强） ==========
                if file_ext == '.pdf':
                    # PDF三层处理策略
                    pdf_data = None
                    pdf_raw = None

                    # 第1层: 尝试提取PDF嵌入文本（数字PDF）
                    embedded_text = _extract_pdf_text(file_path)
                    if embedded_text and len(embedded_text) >= 20:
                        print(f"[PDF] 提取到嵌入文本({len(embedded_text)}字符)，使用LLM结构化")
                        pdf_data, pdf_raw = extract_from_ocr_text(embedded_text, ai_prompt)
                        if 'error' in pdf_data:
                            print(f"[PDF] 嵌入文本结构化失败，尝试OCR")
                            pdf_data = None

                    # 第2层: 尝试本地OCR（扫描件PDF）
                    if pdf_data is None and HAS_TESSERACT:
                        try:
                            ocr_text = local_ocr_pdf(file_path)
                            if ocr_text and len(ocr_text) >= 10:
                                print(f"[PDF] OCR识别成功({len(ocr_text)}字符)，使用LLM结构化")
                                pdf_data, pdf_raw = extract_from_ocr_text(ocr_text, ai_prompt)
                                if 'error' in pdf_data:
                                    print(f"[PDF] OCR文本结构化失败，回退到多模态识别")
                                    pdf_data = None
                        except Exception as e:
                            print(f"[PDF] OCR处理失败: {e}")

                    # 第3层: 回退到多模态逐页识别
                    if pdf_data is None:
                        print(f"[PDF] 使用多模态模型逐页识别")
                        image_paths = pdf_to_images(file_path)
                        for img_path in image_paths:
                            try:
                                data, raw_text = extract_medical_data_multimodal(img_path, ai_prompt)
                                if "error" in data:
                                    errors.append(f"{file.filename}: {data['error']}")
                                    continue
                                case_number = f"CASE_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
                                record_id = str(uuid.uuid4())
                                create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                                confidence_data = data.pop('confidence', {})
                                conn = get_db()
                                c = conn.cursor()
                                c.execute('''INSERT INTO medical_records
                                    (id, case_number, original_filename, role_id, template_id,
                                     extracted_data, confidence_data, raw_text, create_time, source_type, module_type)
                                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'image', ?)''',
                                    (record_id, case_number, file.filename, role_id, template_id,
                                     json.dumps(data, ensure_ascii=False),
                                     json.dumps(confidence_data, ensure_ascii=False),
                                     raw_text, create_time, module_type or 'image_ocr'))
                                conn.commit()
                                conn.close()
                                results.append({
                                    "id": record_id, "case_number": case_number,
                                    "filename": file.filename, "role_id": role_id,
                                    "template_name": template_name, "display_layout": display_layout,
                                    "source_type": "image", "module_type": module_type or "image_ocr",
                                    "data": data, "confidence": confidence_data, "create_time": create_time
                                })
                            finally:
                                try:
                                    if os.path.exists(img_path):
                                        os.remove(img_path)
                                except Exception:
                                    pass
                        continue  # PDF逐页处理已完成，跳过后续单记录逻辑

                    # PDF前两层成功，保存单条记录
                    data = pdf_data
                    raw_text = pdf_raw
                    if "error" in data:
                        errors.append(f"{file.filename}: {data['error']}")
                        continue
                    case_number = f"CASE_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
                    record_id = str(uuid.uuid4())
                    create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    confidence_data = data.pop('confidence', {})
                    conn = get_db()
                    c = conn.cursor()
                    c.execute('''INSERT INTO medical_records
                        (id, case_number, original_filename, role_id, template_id,
                         extracted_data, confidence_data, raw_text, create_time, source_type, module_type)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'image', ?)''',
                        (record_id, case_number, file.filename, role_id, template_id,
                         json.dumps(data, ensure_ascii=False),
                         json.dumps(confidence_data, ensure_ascii=False),
                         raw_text, create_time, module_type or 'image_ocr'))
                    conn.commit()
                    conn.close()
                    results.append({
                        "id": record_id, "case_number": case_number,
                        "filename": file.filename, "role_id": role_id,
                        "template_name": template_name, "display_layout": display_layout,
                        "source_type": "image", "module_type": module_type or "image_ocr",
                        "data": data, "confidence": confidence_data, "create_time": create_time
                    })
                else:
                    # ========== 单张图片处理（使用调度器） ==========
                    data, raw_text = extract_medical_data(file_path, ai_prompt)

                    if "error" in data:
                        errors.append(f"{file.filename}: {data['error']}")
                        continue

                    case_number = f"CASE_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
                    record_id = str(uuid.uuid4())
                    create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

                    confidence_data = data.pop('confidence', {})

                    conn = get_db()
                    c = conn.cursor()
                    c.execute('''INSERT INTO medical_records
                        (id, case_number, original_filename, role_id, template_id,
                         extracted_data, confidence_data, raw_text, create_time, source_type, module_type)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'image', ?)''',
                        (record_id, case_number, file.filename, role_id, template_id,
                         json.dumps(data, ensure_ascii=False),
                         json.dumps(confidence_data, ensure_ascii=False),
                         raw_text, create_time, module_type or 'image_ocr'))
                    conn.commit()
                    conn.close()

                    results.append({
                        "id": record_id,
                        "case_number": case_number,
                        "filename": file.filename,
                        "role_id": role_id,
                        "template_name": template_name,
                        "display_layout": display_layout,
                        "source_type": "image",
                        "module_type": module_type or "image_ocr",
                        "data": data,
                        "confidence": confidence_data,
                        "create_time": create_time
                    })

        except Exception as e:
            errors.append(f"{file.filename}: 识别失败 - {str(e)}")
        finally:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception:
                pass

    return jsonify({
        "status": "success" if results else "error",
        "results": results,
        "errors": errors,
        "msg": f"成功识别 {len(results)} 份" + (f"，{len(errors)} 份失败" if errors else "")
    })


def _process_audio_file(audio_path, filename, role_id, template_id,
                        ai_prompt, template_name, display_layout):
    """处理单个音频文件：语音转写 → 结构化提取 → 可选质性分析"""
    try:
        # 1. 语音转文字
        transcript_result = transcribe_audio(audio_path)
        transcript_text = transcript_result['text']

        # 1.5 脱敏：在发送远程LLM之前，对转录文本进行敏感信息脱敏
        transcript_text, _privacy_report = desensitize_text(transcript_text)

        # 2. 用AI从转录文本提取结构化数据
        data, raw_text = extract_from_transcript(transcript_text, ai_prompt)

        if "error" in data:
            return {"error": data.get('error', '文本提取失败')}

        # 3. 质性分析（仅科研角色）
        qual_result = None
        if role_id == 'other':
            try:
                qual_result = qualitative_analysis(transcript_text)
            except Exception as e:
                print(f"[WARN] 质性分析失败: {e}")

        # 4. 生成记录
        case_number = f"AUDIO_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
        record_id = str(uuid.uuid4())
        create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        confidence_data = data.pop('confidence', {})

        # 5. 存储到数据库
        conn = get_db()
        c = conn.cursor()
        c.execute('''INSERT INTO medical_records
            (id, case_number, original_filename, role_id, template_id,
             extracted_data, confidence_data, raw_text, create_time,
             source_type, audio_transcript, qualitative_data, module_type)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'audio', ?, ?, 'voice_input')''',
            (record_id, case_number, filename, role_id, template_id,
             json.dumps(data, ensure_ascii=False),
             json.dumps(confidence_data, ensure_ascii=False),
             raw_text, create_time,
             transcript_text,
             json.dumps(qual_result, ensure_ascii=False) if qual_result else None))
        conn.commit()
        conn.close()

        return {
            "id": record_id,
            "case_number": case_number,
            "filename": filename,
            "role_id": role_id,
            "template_name": template_name,
            "display_layout": display_layout,
            "source_type": "audio",
            "transcript": transcript_text,
            "qualitative_analysis": qual_result,
            "data": data,
            "confidence": confidence_data,
            "create_time": create_time
        }
    except Exception as e:
        return {"error": f"音频处理失败: {str(e)}"}


@app.route('/upload_text', methods=['POST'])
def upload_text():
    """文本输入模块：处理粘贴文本或txt/docx文件上传"""
    role_id = request.form.get('role_id', 'general')
    template_id = request.form.get('template_id', 'tpl_researcher_default')
    text_content = request.form.get('text_content', '').strip()

    # 查询模板
    conn = get_db()
    c = conn.cursor()
    c.execute("SELECT ai_prompt, template_name, display_layout FROM extraction_templates WHERE template_id=?",
              (template_id,))
    tpl_row = c.fetchone()
    conn.close()

    if not tpl_row:
        return jsonify({"status": "error", "msg": "模板不存在"})

    ai_prompt = tpl_row['ai_prompt']
    template_name = tpl_row['template_name']
    display_layout = tpl_row['display_layout']

    results = []
    errors = []

    # 模式1：直接粘贴文本
    if text_content:
        try:
            processed_text = _preprocess_text(text_content)
            if len(processed_text) < 5:
                return jsonify({"status": "error", "msg": "文本内容过短，请输入更多内容"})

            # 脱敏：在发送远程LLM之前，对文本进行敏感信息脱敏
            processed_text, _privacy_report = desensitize_text(processed_text)

            data, raw_text = extract_from_transcript(processed_text, ai_prompt)
            if "error" in data:
                return jsonify({"status": "error", "msg": data.get('error', '文本提取失败')})

            case_number = f"TEXT_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
            record_id = str(uuid.uuid4())
            create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            confidence_data = data.pop('confidence', {})

            conn = get_db()
            c = conn.cursor()
            c.execute('''INSERT INTO medical_records
                (id, case_number, original_filename, role_id, template_id,
                 extracted_data, confidence_data, raw_text, create_time,
                 source_type, module_type, text_source)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'text', 'text_input', ?)''',
                (record_id, case_number, '粘贴文本', role_id, template_id,
                 json.dumps(data, ensure_ascii=False),
                 json.dumps(confidence_data, ensure_ascii=False),
                 raw_text, create_time, processed_text))
            conn.commit()
            conn.close()

            results.append({
                "id": record_id,
                "case_number": case_number,
                "filename": "粘贴文本",
                "role_id": role_id,
                "template_name": template_name,
                "display_layout": display_layout,
                "source_type": "text",
                "module_type": "text_input",
                "text_source": processed_text,
                "data": data,
                "confidence": confidence_data,
                "create_time": create_time
            })
        except Exception as e:
            errors.append(f"文本处理失败: {str(e)}")

    # 模式2：文件上传
    files = request.files.getlist('files')
    for file in files:
        if not file or file.filename == '':
            continue
        if not is_text_file(file.filename):
            errors.append(f"不支持的格式: {file.filename}")
            continue

        file_ext = os.path.splitext(file.filename)[1].lower()
        temp_name = f"{uuid.uuid4().hex}{file_ext}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_name)
        file.save(file_path)

        try:
            raw_file_text = _parse_text_file(file_path)
            processed_text = _preprocess_text(raw_file_text)
            if len(processed_text) < 5:
                errors.append(f"{file.filename}: 文件内容过短或为空")
                continue

            # 脱敏：在发送远程LLM之前，对文本进行敏感信息脱敏
            processed_text, _privacy_report = desensitize_text(processed_text)

            data, raw_text = extract_from_transcript(processed_text, ai_prompt)
            if "error" in data:
                errors.append(f"{file.filename}: {data.get('error', '提取失败')}")
                continue

            case_number = f"TEXT_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
            record_id = str(uuid.uuid4())
            create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            confidence_data = data.pop('confidence', {})

            conn = get_db()
            c = conn.cursor()
            c.execute('''INSERT INTO medical_records
                (id, case_number, original_filename, role_id, template_id,
                 extracted_data, confidence_data, raw_text, create_time,
                 source_type, module_type, text_source)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'text', 'text_input', ?)''',
                (record_id, case_number, file.filename, role_id, template_id,
                 json.dumps(data, ensure_ascii=False),
                 json.dumps(confidence_data, ensure_ascii=False),
                 raw_text, create_time, processed_text))
            conn.commit()
            conn.close()

            results.append({
                "id": record_id,
                "case_number": case_number,
                "filename": file.filename,
                "role_id": role_id,
                "template_name": template_name,
                "display_layout": display_layout,
                "source_type": "text",
                "module_type": "text_input",
                "text_source": processed_text,
                "data": data,
                "confidence": confidence_data,
                "create_time": create_time
            })
        except Exception as e:
            errors.append(f"{file.filename}: {str(e)}")
        finally:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception:
                pass

    if not results and not errors:
        return jsonify({"status": "error", "msg": "请输入文本或上传文件"})

    return jsonify({
        "status": "success" if results else "error",
        "results": results,
        "errors": errors,
        "msg": f"成功处理 {len(results)} 份" + (f"，{len(errors)} 份失败" if errors else "")
    })


@app.route('/qualitative_analyze', methods=['POST'])
def qualitative_analyze():
    """质性研究模块：独立的质性分析入口"""
    analysis_type = request.form.get('analysis_type', 'interview')
    text_content = request.form.get('text_content', '').strip()

    results = []
    errors = []
    transcript_text = ''
    source_filename = ''

    # 模式1：音频文件 → 转录 → 分析
    files = request.files.getlist('files')
    for file in files:
        if not file or file.filename == '':
            continue

        file_ext = os.path.splitext(file.filename)[1].lower()
        temp_name = f"{uuid.uuid4().hex}{file_ext}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_name)
        file.save(file_path)

        try:
            if is_audio_file(file.filename):
                transcript_result = transcribe_audio(file_path)
                transcript_text = transcript_result['text']
                source_filename = file.filename
            elif is_text_file(file.filename):
                raw_text = _parse_text_file(file_path)
                transcript_text = _preprocess_text(raw_text)
                source_filename = file.filename
            else:
                errors.append(f"不支持的格式: {file.filename}，请上传音频或文本文件")
                continue

            if len(transcript_text) < 10:
                errors.append(f"{file.filename}: 内容过短，无法进行质性分析")
                continue

            # 脱敏：在发送远程LLM和存储之前，对转录文本进行敏感信息脱敏
            transcript_text, _privacy_report = desensitize_text(transcript_text)

            qual_result = qualitative_analysis_enhanced(transcript_text, analysis_type)

            case_number = f"QUAL_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
            record_id = str(uuid.uuid4())
            create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

            conn = get_db()
            c = conn.cursor()
            c.execute('''INSERT INTO medical_records
                (id, case_number, original_filename, role_id, template_id,
                 extracted_data, raw_text, create_time,
                 source_type, module_type, audio_transcript, qualitative_data, analysis_type)
                VALUES (?, ?, ?, NULL, NULL, NULL, NULL, ?, ?, 'qualitative', ?, ?, ?)''',
                (record_id, case_number, source_filename, create_time,
                 'audio' if is_audio_file(file.filename) else 'text',
                 transcript_text,
                 json.dumps(qual_result, ensure_ascii=False),
                 analysis_type))
            conn.commit()
            conn.close()

            results.append({
                "id": record_id,
                "case_number": case_number,
                "filename": source_filename,
                "source_type": "audio" if is_audio_file(file.filename) else "text",
                "module_type": "qualitative",
                "analysis_type": analysis_type,
                "transcript": transcript_text,
                "qualitative_analysis": qual_result,
                "create_time": create_time
            })
        except Exception as e:
            errors.append(f"{file.filename}: {str(e)}")
        finally:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception:
                pass

    # 模式2：直接粘贴文本
    if text_content and not results:
        try:
            processed_text = _preprocess_text(text_content)
            if len(processed_text) < 10:
                return jsonify({"status": "error", "msg": "文本内容过短，无法进行质性分析"})

            # 脱敏：在发送远程LLM和存储之前，对文本进行敏感信息脱敏
            processed_text, _privacy_report = desensitize_text(processed_text)

            qual_result = qualitative_analysis_enhanced(processed_text, analysis_type)

            case_number = f"QUAL_{datetime.now().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6].upper()}"
            record_id = str(uuid.uuid4())
            create_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

            conn = get_db()
            c = conn.cursor()
            c.execute('''INSERT INTO medical_records
                (id, case_number, original_filename, role_id, template_id,
                 extracted_data, raw_text, create_time,
                 source_type, module_type, audio_transcript, qualitative_data, analysis_type)
                VALUES (?, ?, ?, NULL, NULL, NULL, NULL, ?, 'text', 'qualitative', ?, ?, ?)''',
                (record_id, case_number, '粘贴文本', create_time,
                 processed_text,
                 json.dumps(qual_result, ensure_ascii=False),
                 analysis_type))
            conn.commit()
            conn.close()

            results.append({
                "id": record_id,
                "case_number": case_number,
                "filename": "粘贴文本",
                "source_type": "text",
                "module_type": "qualitative",
                "analysis_type": analysis_type,
                "transcript": processed_text,
                "qualitative_analysis": qual_result,
                "create_time": create_time
            })
        except Exception as e:
            errors.append(f"质性分析失败: {str(e)}")

    if not results and not errors:
        return jsonify({"status": "error", "msg": "请上传文件或输入文本"})

    return jsonify({
        "status": "success" if results else "error",
        "results": results,
        "errors": errors,
        "msg": f"成功分析 {len(results)} 份" + (f"，{len(errors)} 份失败" if errors else "")
    })


@app.route('/records', methods=['GET'])
def get_records():
    role_id = request.args.get('role_id', None)
    module_type = request.args.get('module_type', None)
    conn = get_db()
    c = conn.cursor()

    conditions = []
    params = []
    if role_id:
        conditions.append("role_id=?")
        params.append(role_id)
    if module_type:
        conditions.append("module_type=?")
        params.append(module_type)

    where_clause = (" WHERE " + " AND ".join(conditions)) if conditions else ""
    c.execute(f'''SELECT id, case_number, original_filename, role_id, template_id,
        create_time, source_type, module_type
        FROM medical_records{where_clause} ORDER BY create_time DESC''', params)
    rows = c.fetchall()
    conn.close()

    records = []
    for row in rows:
        records.append({
            "id": row['id'],
            "case_number": row['case_number'],
            "filename": row['original_filename'],
            "role_id": row['role_id'] or 'general',
            "template_id": row['template_id'] or '',
            "create_time": row['create_time'],
            "source_type": row['source_type'] or 'image',
            "module_type": row['module_type'] or 'image_ocr'
        })
    return jsonify({"status": "success", "records": records})


@app.route('/record/<record_id>', methods=['GET'])
def get_record_detail(record_id):
    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT * FROM medical_records WHERE id = ?', (record_id,))
    row = c.fetchone()

    # 查模板信息
    template_name = ''
    display_layout = 'table'
    if row and row['template_id']:
        c.execute("SELECT template_name, display_layout FROM extraction_templates WHERE template_id=?",
                  (row['template_id'],))
        tpl = c.fetchone()
        if tpl:
            template_name = tpl['template_name']
            display_layout = tpl['display_layout']
    conn.close()

    if not row:
        return jsonify({"status": "error", "msg": "记录不存在"})

    role_id = row['role_id'] or 'other'

    # 兼容旧数据
    if row['extracted_data']:
        extracted_data = json.loads(row['extracted_data'])
        confidence_data = json.loads(row['confidence_data']) if row['confidence_data'] else {}
    else:
        # 旧格式兼容
        extracted_data = {
            'demographics': json.loads(row['demographics']) if row['demographics'] else {},
            'lab_tests': json.loads(row['lab_tests']) if row['lab_tests'] else [],
            'treatment': json.loads(row['treatment']) if row['treatment'] else {},
        }
        confidence_data = json.loads(row['confidence']) if row['confidence'] else {}
        display_layout = 'table'
        template_name = '综合科研数据提取'

    return jsonify({
        "status": "success",
        "data": {
            "id": row['id'],
            "case_number": row['case_number'],
            "filename": row['original_filename'],
            "role_id": role_id,
            "template_name": template_name,
            "display_layout": display_layout,
            "extracted_data": extracted_data,
            "confidence": confidence_data,
            "create_time": row['create_time'],
            "source_type": row['source_type'] or 'image',
            "module_type": row['module_type'] or 'image_ocr',
            "audio_transcript": row['audio_transcript'] if (row['source_type'] == 'audio') else None,
            "qualitative_data": json.loads(row['qualitative_data']) if row['qualitative_data'] else None,
            "text_source": row['text_source'] if (row['source_type'] == 'text') else None,
            "analysis_type": row['analysis_type'] if row['module_type'] == 'qualitative' else None
        }
    })


@app.route('/record/<record_id>', methods=['PUT'])
def update_record(record_id):
    update_data = request.get_json()
    if not update_data:
        return jsonify({"status": "error", "msg": "无更新数据"})

    conn = get_db()
    c = conn.cursor()
    if 'extracted_data' in update_data:
        c.execute("UPDATE medical_records SET extracted_data=? WHERE id=?",
                  (json.dumps(update_data['extracted_data'], ensure_ascii=False), record_id))
    if 'confidence' in update_data:
        c.execute("UPDATE medical_records SET confidence_data=? WHERE id=?",
                  (json.dumps(update_data['confidence'], ensure_ascii=False), record_id))
    conn.commit()
    conn.close()
    return jsonify({"status": "success", "msg": "数据已更新"})


@app.route('/record/<record_id>', methods=['DELETE'])
def delete_record(record_id):
    conn = get_db()
    c = conn.cursor()
    c.execute('DELETE FROM medical_records WHERE id = ?', (record_id,))
    conn.commit()
    conn.close()
    return jsonify({"status": "success", "msg": "记录已删除"})


@app.route('/export', methods=['POST'])
def export_excel():
    req_data = request.get_json() or {}
    record_ids = req_data.get('record_ids', [])
    role_filter = req_data.get('role_id', None)

    conn = get_db()
    c = conn.cursor()

    if record_ids:
        placeholders = ','.join(['?'] * len(record_ids))
        c.execute(f'SELECT * FROM medical_records WHERE id IN ({placeholders})', record_ids)
    elif role_filter:
        c.execute('SELECT * FROM medical_records WHERE role_id=? ORDER BY create_time', (role_filter,))
    else:
        c.execute('SELECT * FROM medical_records ORDER BY create_time')

    rows = c.fetchall()
    conn.close()

    if not rows:
        return jsonify({"status": "error", "msg": "暂无可导出的数据"})

    data_list = _rows_to_export_list(rows)
    excel_path = generate_excel(data_list)
    return send_file(excel_path, as_attachment=True,
                     download_name=f"临床数据_{datetime.now().strftime('%Y%m%d')}.xlsx")


@app.route('/export_all', methods=['GET'])
def export_all_excel():
    role_filter = request.args.get('role_id', None)
    conn = get_db()
    c = conn.cursor()
    if role_filter:
        c.execute('SELECT * FROM medical_records WHERE role_id=? ORDER BY create_time', (role_filter,))
    else:
        c.execute('SELECT * FROM medical_records ORDER BY create_time')
    rows = c.fetchall()
    conn.close()

    if not rows:
        return jsonify({"status": "error", "msg": "暂无可导出的数据"})

    data_list = _rows_to_export_list(rows)
    excel_path = generate_excel(data_list)
    return send_file(excel_path, as_attachment=True,
                     download_name=f"临床数据_{datetime.now().strftime('%Y%m%d')}.xlsx")


def _rows_to_export_list(rows):
    """将数据库行转为导出数据列表"""
    conn = get_db()
    c = conn.cursor()
    data_list = []
    for row in rows:
        role_id = row['role_id'] or 'other'
        template_name = ''
        if row['template_id']:
            c.execute("SELECT template_name FROM extraction_templates WHERE template_id=?",
                      (row['template_id'],))
            tpl = c.fetchone()
            if tpl:
                template_name = tpl['template_name']

        if row['extracted_data']:
            extracted = json.loads(row['extracted_data'])
            conf = json.loads(row['confidence_data']) if row['confidence_data'] else {}
        else:
            extracted = {
                'demographics': json.loads(row['demographics']) if row['demographics'] else {},
                'lab_tests': json.loads(row['lab_tests']) if row['lab_tests'] else [],
                'treatment': json.loads(row['treatment']) if row['treatment'] else {},
            }
            conf = json.loads(row['confidence']) if row['confidence'] else {}

        data_list.append({
            'case_number': row['case_number'],
            'create_time': row['create_time'],
            'role_id': role_id,
            'template_name': template_name,
            'extracted_data': extracted,
            'confidence_data': conf,
            'source_type': row['source_type'] or 'image',
            'audio_transcript': row['audio_transcript'] if row['source_type'] == 'audio' else None,
            'qualitative_data': json.loads(row['qualitative_data']) if row['qualitative_data'] else None,
        })
    conn.close()
    return data_list


@app.route('/clean', methods=['POST'])
def clean_all():
    try:
        shutil.rmtree(app.config['UPLOAD_FOLDER'])
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    except Exception:
        pass
    conn = get_db()
    c = conn.cursor()
    c.execute('DELETE FROM medical_records')
    conn.commit()
    conn.close()
    return jsonify({"status": "success", "msg": "所有数据已清理"})


@app.route('/stats', methods=['GET'])
def get_stats():
    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT COUNT(*) as cnt FROM medical_records')
    total = c.fetchone()['cnt']
    # 按角色统计
    c.execute("SELECT role_id, COUNT(*) as cnt FROM medical_records GROUP BY role_id")
    by_role = {}
    for row in c.fetchall():
        by_role[row['role_id'] or 'other'] = row['cnt']
    conn.close()
    return jsonify({"status": "success", "total_records": total, "by_role": by_role})


@app.route('/data_analysis/fields', methods=['GET'])
def data_analysis_fields():
    """获取选中记录的所有可用数值字段"""
    record_ids = request.args.getlist('ids')
    if not record_ids:
        return jsonify({"status": "error", "msg": "请选择记录"})

    conn = get_db()
    c = conn.cursor()
    placeholders = ','.join(['?'] * len(record_ids))
    c.execute(f'SELECT extracted_data FROM medical_records WHERE id IN ({placeholders})', record_ids)
    rows = c.fetchall()
    conn.close()

    all_paths = set()
    for row in rows:
        if row['extracted_data']:
            data = json.loads(row['extracted_data'])
            paths = _collect_field_paths(data)
            all_paths.update(paths)

    # 过滤出含数值的字段
    numeric_fields = []
    text_fields = []
    for path in sorted(all_paths):
        has_numeric = False
        for row in rows:
            if row['extracted_data']:
                data = json.loads(row['extracted_data'])
                val = _extract_nested_field(data, path)
                if _is_numeric(val):
                    has_numeric = True
                    break
        if has_numeric:
            numeric_fields.append(path)
        else:
            text_fields.append(path)

    return jsonify({
        "status": "success",
        "numeric_fields": numeric_fields,
        "text_fields": text_fields
    })


@app.route('/data_analysis/analyze', methods=['POST'])
def data_analysis_analyze():
    """数据分析模块：对选中记录进行统计分析"""
    req_data = request.get_json()
    if not req_data:
        return jsonify({"status": "error", "msg": "无请求数据"})

    record_ids = req_data.get('record_ids', [])
    fields = req_data.get('fields', [])
    analysis_type = req_data.get('analysis_type', 'descriptive')

    if not record_ids:
        return jsonify({"status": "error", "msg": "请选择至少1条记录"})
    if not fields:
        return jsonify({"status": "error", "msg": "请选择至少1个分析字段"})

    try:
        result = analyze_structured_data(record_ids, fields, analysis_type)
        return jsonify({"status": "success", **result})
    except Exception as e:
        return jsonify({"status": "error", "msg": f"分析失败: {str(e)}"})


# ========== 启动入口 ==========
if __name__ == '__main__':
    import sys
    port = int(sys.argv[1]) if len(sys.argv) > 1 else 7860
    print("=" * 50)
    print("  临床科研病历AI识别与结构化提取工具（多角色版）")
    print(f"  访问地址: http://localhost:{port}")
    print("  按 Ctrl+C 停止服务")
    print("=" * 50)
    init_db()
    app.run(host='0.0.0.0', port=port, debug=False)
